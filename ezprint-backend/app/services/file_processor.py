"""PDF introspection + color classification.

Light-weight port of the parts of `shared/file_processor.py` needed by the
backend. Heavy rendering / SumatraPDF printing stays on the Windows agent
side, where it belongs. This module only inspects files pulled back from
MinIO to fill in page counts and detect color pages for billing.

The goal is best-effort; if a file is corrupt or a library is missing we
return safe defaults (1 page, 0 color pages) rather than crashing a request.
"""
from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class PageStats:
    total_pages: int
    color_pages: int


def _classify_pdf_with_fitz(data: bytes) -> PageStats:
    try:
        import fitz  # PyMuPDF
    except Exception:
        raise

    total = 0
    color_pages = 0
    with fitz.open(stream=data, filetype="pdf") as doc:
        total = doc.page_count
        for page in doc:
            # Quick test: render at low resolution and look for non-gray pixels.
            pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5), alpha=False)
            if pix.n >= 3:
                sample = pix.samples
                stride = max(1, len(sample) // 3_000)  # sample up to ~1000 pixels
                has_color = False
                for i in range(0, len(sample) - 2, 3 * stride):
                    r, g, b = sample[i], sample[i + 1], sample[i + 2]
                    if abs(r - g) > 8 or abs(g - b) > 8 or abs(r - b) > 8:
                        has_color = True
                        break
                if has_color:
                    color_pages += 1
    return PageStats(total_pages=total, color_pages=color_pages)


def _classify_pdf_with_pypdf(data: bytes) -> PageStats:
    import PyPDF2  # type: ignore

    reader = PyPDF2.PdfReader(io.BytesIO(data))
    return PageStats(total_pages=len(reader.pages), color_pages=0)


def classify_pdf(data: bytes) -> PageStats:
    """Return (total_pages, color_pages). Best-effort; never raises."""
    try:
        return _classify_pdf_with_fitz(data)
    except Exception as exc:
        logger.info("fitz classification unavailable (%s); falling back to PyPDF2", exc)
    try:
        return _classify_pdf_with_pypdf(data)
    except Exception as exc:
        logger.warning("pypdf2 fallback failed: %s", exc)
    return PageStats(total_pages=1, color_pages=0)


def classify_image(data: bytes) -> PageStats:
    """Single image = 1 page; call it color if mode has >1 band."""
    try:
        from PIL import Image
    except Exception:
        return PageStats(total_pages=1, color_pages=0)
    try:
        with Image.open(io.BytesIO(data)) as img:
            is_color = img.mode not in {"1", "L", "P"}  # crude but good enough
            return PageStats(total_pages=1, color_pages=1 if is_color else 0)
    except Exception as exc:
        logger.warning("image classify failed: %s", exc)
        return PageStats(total_pages=1, color_pages=0)


def _estimate_docx_page_count(data: bytes) -> int:
    """Rough page count for .docx (print-style estimate; matches browser slices ~OK)."""
    try:
        from docx import Document
    except Exception as exc:
        logger.info("python-docx unavailable: %s", exc)
        return 1
    try:
        doc = Document(io.BytesIO(data))
        words = 0
        for p in doc.paragraphs:
            t = p.text or ""
            if t.strip():
                words += len(re.findall(r"[\w']+", t))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text or ""
                    if t.strip():
                        words += len(re.findall(r"[\w']+", t))
    except Exception as exc:
        logger.info("docx page estimate failed: %s", exc)
        return 1
    # ~250 words / A4 page (body text; aligns loosely with in-browser A4-height slices)
    wpp = 250
    return max(1, min(500, (max(words, 1) + wpp - 1) // wpp))


def classify_docx(data: bytes) -> PageStats:
    n = _estimate_docx_page_count(data)
    return PageStats(total_pages=n, color_pages=0)


def classify_bytes(file_type: str, data: bytes) -> PageStats:
    ft = (file_type or "").lower().lstrip(".")
    if ft == "pdf":
        return classify_pdf(data)
    if ft in {"jpg", "jpeg", "png", "bmp", "tiff", "webp", "gif"}:
        return classify_image(data)
    if ft == "docx":
        return classify_docx(data)
    # For xlsx / legacy .doc / etc. we just stamp a single page and bill BW by default.
    return PageStats(total_pages=1, color_pages=0)


def _pdf_num_pages(data: bytes) -> int:
    try:
        import fitz  # PyMuPDF

        with fitz.open(stream=data, filetype="pdf") as doc:
            return int(doc.page_count) or 1
    except Exception:
        pass
    try:
        import PyPDF2  # type: ignore

        return max(1, len(PyPDF2.PdfReader(io.BytesIO(data)).pages))
    except Exception:
        return 1


def parse_page_range_to_indices(page_range: str | None, num_pages: int) -> list[int]:
    """1-based page indices to bill, matching the customer upload UI.

    Empty / whitespace `page_range` means all pages. Invalid segments are skipped; if
    nothing parses, all pages are used (same as the browser helper).
    """
    if not num_pages or num_pages < 1:
        return []
    s = (page_range or "").strip()
    if not s:
        return list(range(1, num_pages + 1))
    raw = "".join(s.split())
    out: list[int] = []
    seen: set[int] = set()
    for part in raw.split(","):
        if not part:
            continue
        if "-" in part:
            ab = part.split("-", 1)
            try:
                lo = max(1, int(ab[0]))
                hi = min(num_pages, int(ab[1]))
            except ValueError:
                continue
            for p in range(lo, hi + 1):
                if 1 <= p <= num_pages and p not in seen:
                    seen.add(p)
                    out.append(p)
        else:
            try:
                p = int(part)
            except ValueError:
                continue
            if 1 <= p <= num_pages and p not in seen:
                seen.add(p)
                out.append(p)
    if not out:
        return list(range(1, num_pages + 1))
    return sorted(out)


def _classify_pdf_pages_fitz(data: bytes, one_based_indices: list[int]) -> PageStats:
    import fitz  # PyMuPDF

    color_pages = 0
    with fitz.open(stream=data, filetype="pdf") as doc:
        n = doc.page_count
        valid = [p for p in sorted(set(one_based_indices)) if 1 <= p <= n]
        for p in valid:
            page = doc[p - 1]
            pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5), alpha=False)
            if pix.n >= 3:
                sample = pix.samples
                stride = max(1, len(sample) // 3_000)
                has_color = False
                for i in range(0, len(sample) - 2, 3 * stride):
                    r, g, b = sample[i], sample[i + 1], sample[i + 2]
                    if abs(r - g) > 8 or abs(g - b) > 8 or abs(r - b) > 8:
                        has_color = True
                        break
                if has_color:
                    color_pages += 1
    return PageStats(total_pages=len(valid), color_pages=color_pages)


def classify_pdf_for_page_range(data: bytes, page_range: str | None) -> PageStats:
    """Classify only pages included in *page_range* (for billing after finalize)."""
    n = _pdf_num_pages(data)
    indices = parse_page_range_to_indices(page_range, n)
    if not indices:
        return PageStats(total_pages=1, color_pages=0)
    if len(indices) == n and indices[0] == 1 and indices[-1] == n and len(set(indices)) == n:
        return classify_pdf(data)
    try:
        return _classify_pdf_pages_fitz(data, indices)
    except Exception as exc:
        logger.info("fitz page-range classification unavailable (%s); using page count only", exc)
    valid = [p for p in indices if 1 <= p <= n]
    return PageStats(total_pages=max(1, len(valid)), color_pages=0)


def classify_bytes_for_job(file_type: str, data: bytes, page_range: str | None) -> PageStats:
    """Like ``classify_bytes`` but applies PDF *page_range* to page/color counts for billing."""
    ft = (file_type or "").lower().lstrip(".")
    if ft == "pdf":
        return classify_pdf_for_page_range(data, page_range)
    if ft == "docx":
        n = _estimate_docx_page_count(data)
        idx = parse_page_range_to_indices(page_range, n)
        if not idx:
            return PageStats(total_pages=1, color_pages=0)
        return PageStats(total_pages=len(idx), color_pages=0)
    return classify_bytes(file_type, data)
