"""
File processing utilities for document handling
"""
import os
import uuid
import io
import subprocess
import tempfile
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
import PyPDF2
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4, A3, legal
from reportlab.lib.utils import ImageReader
from shared.config import UPLOAD_FOLDER, ALLOWED_EXTENSIONS
import re  # added for shop_id sanitization
import time
import requests
from urllib.parse import urlparse

import logging
logger = logging.getLogger(__name__)
def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# --- URL Handling Utilities ---
def is_url(path):
    """Check if a path is a URL"""
    if not isinstance(path, str):
        return False
    return path.startswith(('http://', 'https://'))

def ensure_local_path(file_path):
    """
    Ensures that file_path is a local path. 
    If it's a URL, downloads it to a temp file.
    Returns: (local_path, is_temp)
    """
    if not is_url(file_path):
        return file_path, False
        
    try:
        # Download file (using plain GET as these are public or signed Cloudinary URLs)
        response = requests.get(file_path, stream=True, timeout=30, headers={})
        response.raise_for_status()
        
        # Determine extension from URL or content type
        parsed = urlparse(file_path)
        ext = os.path.splitext(parsed.path)[1].lower()
        if not ext:
            content_type = response.headers.get('content-type', '').lower()
            if 'pdf' in content_type: ext = '.pdf'
            elif 'word' in content_type or 'officedocument' in content_type: ext = '.docx'
            elif 'image' in content_type: ext = '.png'
            else: ext = '.pdf' # Default to PDF
            
        fd, temp_path = tempfile.mkstemp(suffix=ext)
        os.close(fd)
        
        with open(temp_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
                
        return temp_path, True
    except Exception as e:
        print(f"Error downloading from URL {file_path}: {e}")
        return file_path, False

def _compress_file_for_upload(file_path, file_extension, target_size_bytes=9 * 1024 * 1024):
    """
    Silently compress file if too large for Cloudinary.
    Returns compressed file path (may be same as input if compression not needed/possible).
    """
    try:
        current_size = os.path.getsize(file_path)
        if current_size <= target_size_bytes:
            return file_path  # No compression needed

        logger.info(f"File too large ({current_size} bytes), compressing...")

        if file_extension in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'gif']:
            # Compress image
            img = Image.open(file_path)
            img.load()  # Force load to catch corrupt images early
            compressed_path = file_path + '_compressed.jpg'

            # Convert to RGB with white background for transparency
            if img.mode in ('RGBA', 'LA'):
                bg = Image.new("RGB", img.size, (255, 255, 255))
                bg.paste(img, mask=img.split()[-1])
                img = bg
            elif img.mode in ('P',):
                img = img.convert('RGB')

            # Try progressively lower quality until under limit
            for quality in [85, 70, 55, 40, 25]:
                img.save(compressed_path, 'JPEG', quality=quality, optimize=True)
                if os.path.getsize(compressed_path) <= target_size_bytes:
                    logger.info(f"Compressed to {os.path.getsize(compressed_path)} bytes at quality={quality}")
                    return compressed_path

            # If still too large, resize image with max attempts
            w, h = img.size
            max_attempts = 10
            attempt = 0
            while os.path.getsize(compressed_path) > target_size_bytes and w > 500 and attempt < max_attempts:
                w = int(w * 0.8)
                h = int(h * 0.8)
                img_resized = img.resize((w, h), Image.LANCZOS)
                img_resized.save(compressed_path, 'JPEG', quality=40, optimize=True)
                attempt += 1

            return compressed_path

        elif file_extension == 'pdf':
            # For PDF — try using PyMuPDF to reduce image quality inside PDF
            try:
                import fitz
                doc = fitz.open(file_path)
                compressed_path = file_path + '_compressed.pdf'
                # Save with reduced image quality
                doc.save(compressed_path, deflate=True, garbage=4, clean=True)
                doc.close()
                if os.path.getsize(compressed_path) <= target_size_bytes:
                    logger.info(f"PDF compressed to {os.path.getsize(compressed_path)} bytes")
                    return compressed_path
            except Exception as e:
                logger.warning(f"PDF compression failed: {e}")

        # Could not compress — return original
        return file_path

    except Exception as e:
        logger.warning(f"Compression failed (non-fatal): {e}")
        return file_path  # Return original on error


def save_uploaded_file(file, shop_id):
    """
    Save uploaded file to Cloudinary cloud storage.
    Layer 1: Auto-compress if file too large
    Layer 2: Cloudinary upload
    Layer 3: Fallback to local if Cloudinary fails
    """
    if file and allowed_file(file.filename):
        from shared.cloudinary_helper import upload_file_to_cloudinary

        # Generate unique filename
        file_extension = file.filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4()}.{file_extension}"

        # Save to temporary location first
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, unique_filename)
        file.save(temp_path)

        # Get file size
        file_size = os.path.getsize(temp_path)

        # LAYER 1: Auto-compress if file too large (silent, no user friction)
        MAX_CLOUDINARY_SIZE = 9 * 1024 * 1024  # 9MB safe limit
        compressed_path = temp_path
        if file_size > MAX_CLOUDINARY_SIZE:
            compressed_path = _compress_file_for_upload(temp_path, file_extension)
            compressed_size = os.path.getsize(compressed_path)
            reduction = ((file_size - compressed_size) / file_size) * 100
            logger.info(f"File compressed: {file_size} → {compressed_size} bytes ({reduction:.1f}% reduction)")

        try:
            safe_shop_id = re.sub(r'[^A-Za-z0-9_-]', '_', (shop_id or 'unknown'))[:64]

            # LAYER 2: Upload to Cloudinary
            try:
                cloudinary_url_result, public_id = upload_file_to_cloudinary(
                    compressed_path, safe_shop_id, os.path.basename(compressed_path)
                )
                logger.info(f"Cloudinary upload successful")
                return cloudinary_url_result, file.filename, file_size, file_extension, public_id

            except Exception as cloud_err:
                # LAYER 3: Cloudinary failed — fallback to local storage
                logger.warning(f"Cloudinary upload failed ({cloud_err}), falling back to local storage")

                # Save locally as fallback
                local_upload_dir = os.path.join(
                    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                    'uploads', safe_shop_id
                )
                os.makedirs(local_upload_dir, exist_ok=True)
                local_file_path = os.path.join(local_upload_dir, unique_filename)

                # Copy compressed file to local uploads
                import shutil
                shutil.copy2(compressed_path, local_file_path)

                logger.info(f"File saved locally as fallback: {local_file_path}")
                # Return local path as URL (None public_id = local file)
                local_url = f"/uploads/{safe_shop_id}/{unique_filename}"
                logger.info(f"Local fallback URL: {local_url}")
                return local_url, file.filename, file_size, file_extension, None

        finally:
            # Cleanup compressed file first (if different from original)
            if compressed_path != temp_path and os.path.exists(compressed_path):
                try:
                    os.remove(compressed_path)
                except Exception:
                    pass
            # Cleanup original temp file
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except Exception:
                    pass

    return None, None, None, None, None

def get_page_count(file_path, file_type):
    """
    Get page count for different file types.
    Handles both local paths and Cloudinary URLs.
    """
    # ROOT FIX: Ensure we have a local path for processing
    local_path, is_temp = ensure_local_path(file_path)
    
    try:
        if file_type == 'pdf':
            with open(local_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        
        elif file_type in ['docx', 'doc', 'odt', 'pptx', 'ppt', 'xlsx', 'xls']:
            try:
                # Normalization works on the local path
                norm_pdf = normalize_document_for_preview(local_path, file_type)
                if norm_pdf != local_path and os.path.exists(norm_pdf):
                    return get_page_count(norm_pdf, 'pdf')
                
                # Fallback to rough estimate
                if file_type in ['docx', 'doc']:
                    doc = Document(local_path)
                    count = len(doc.paragraphs) // 10 + 1
                    return max(1, count)
            except Exception as e:
                print(f"Error getting normalized page count: {e}")
            return 1
        
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
            return 1
        
        return 1
    except Exception as e:
        print(f"Error getting page count: {e}")
        return 1
    finally:
        # Cleanup temp file if downloaded
        if is_temp and os.path.exists(local_path):
            try: os.remove(local_path)
            except: pass

def create_preview_image_with_layout(file_path, file_type, page_range="1", page_size="A4", orientation="Portrait", layout_pages=1, layout_type="normal", color_mode="Color"):
    """
    Create a preview image for the document with layout arrangement
    Handles both local paths and Cloudinary URLs.
    """
    # ROOT FIX: Ensure local path
    local_path, is_temp = ensure_local_path(file_path)
    try:
        print(f"Creating layout preview: layout_pages={layout_pages}, layout_type={layout_type}, orientation={orientation}")
        
        # Step 0: Normalize document to PDF for high-fidelity previewing
        if file_type != 'pdf':
            normalized_path = normalize_document_for_preview(local_path, file_type)
            if normalized_path != local_path:
                local_path = normalized_path
                file_type = 'pdf'
        
        # Create preview directory
        if is_url(file_path): # Use original file_path to determine if it was a URL
            preview_dir = UPLOAD_FOLDER / "previews"
        else:
            preview_dir = Path(local_path).parent / "previews"
        preview_dir.mkdir(exist_ok=True)
        
        # Create unique filename with layout info
        orientation_suffix = "L" if orientation == "Landscape" else "P"
        page_size_suffix = page_size.replace(" ", "")
        layout_suffix = f"{layout_pages}up" if layout_pages > 1 else "normal"
        preview_filename = f"preview_{orientation_suffix}_{page_size_suffix}_{layout_suffix}_{uuid.uuid4().hex[:8]}.png"
        preview_path = preview_dir / preview_filename
        
        if file_type == 'pdf':
            # For PDF, create layout-aware preview
            try:
                import fitz  # PyMuPDF for better PDF handling
                
                # Open PDF document
                pdf_document = fitz.open(local_path)
                
                # Parse page range
                page_numbers = parse_page_range(page_range, len(pdf_document))
                
                if not page_numbers:
                    page_numbers = [0]  # Default to first page
                
                # Get the first page from the range
                page_index = page_numbers[0] - 1  # Convert to 0-based index
                
                if 0 <= page_index < len(pdf_document):
                    page = pdf_document[page_index]
                    
                    # Create layout preview based on layout_pages
                    if layout_pages == 1:
                        preview_img = create_single_page_preview(page, page_size, orientation)
                    elif layout_pages in [2,4]:
                        # Keep existing tuned implementations for 2 and 4
                        preview_img = create_2up_preview(page, page_size, orientation) if layout_pages == 2 else create_4up_preview(page, page_size, orientation)
                    else:
                        # Generic N-up using grid tiling
                        from PIL import Image as PILImage
                        # Render a single page image first
                        import fitz
                        page_rect = page.rect
                        zoom_factor = get_zoom_for_page_size(page_size, page_rect) * 0.5
                        if orientation == "Landscape":
                            if page_rect.width < page_rect.height:
                                mat = fitz.Matrix(0, 1, -1, 0, page_rect.height, 0) * fitz.Matrix(zoom_factor, zoom_factor)
                            else:
                                mat = fitz.Matrix(zoom_factor, zoom_factor)
                        else:
                            if page_rect.width > page_rect.height:
                                mat = fitz.Matrix(0, -1, 1, 0, 0, page_rect.width) * fitz.Matrix(zoom_factor, zoom_factor)
                            else:
                                mat = fitz.Matrix(zoom_factor, zoom_factor)
                        pix = page.get_pixmap(matrix=mat, alpha=False)
                        img_data = pix.tobytes("png")
                        single_page_img = PILImage.open(io.BytesIO(img_data))
                        rows, cols = get_grid_for_layout(layout_pages)
                        preview_img = create_generic_grid_preview(single_page_img, rows, cols)
                    
                    # Apply color mode conversion
                    preview_img = apply_color_mode(preview_img, color_mode)
                    
                    # Save the preview image
                    preview_img.save(preview_path, "PNG", quality=95)
                    pdf_document.close()
                    
                    return str(preview_path)
                else:
                    raise Exception(f"Page {page_index + 1} not found in PDF")
                    
            except ImportError:
                # Fallback to basic preview if PyMuPDF not available
                return create_preview_image(local_path, file_type, page_range, page_size, orientation, color_mode)
        
        elif file_type in ['docx', 'doc']:
            # For Word documents, create layout preview
            try:
                doc = Document(local_path)
                text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                if layout_pages == 1:
                    preview_img = create_word_single_page_preview(text_content, page_size, orientation)
                elif layout_pages in [2,4]:
                    preview_img = create_word_2up_preview(text_content, page_size, orientation) if layout_pages == 2 else create_word_4up_preview(text_content, page_size, orientation)
                else:
                    # Generic grid tiling
                    base_img = create_word_single_page_preview(text_content, page_size, orientation)
                    rows, cols = get_grid_for_layout(layout_pages)
                    preview_img = create_generic_grid_preview(base_img, rows, cols)
                
                # Apply color mode conversion
                preview_img = apply_color_mode(preview_img, color_mode)
                
                preview_img.save(preview_path, "PNG", quality=95)
                return str(preview_path)
                
            except Exception as e:
                print(f"Error processing Word document: {e}")
                return create_preview_image(local_path, file_type, page_range, page_size, orientation, color_mode)
        
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
            # For images, create layout preview
            try:
                img = Image.open(local_path)
                
                if layout_pages == 1:
                    preview_img = create_image_single_page_preview(img, page_size, orientation)
                elif layout_pages in [2,4]:
                    preview_img = create_image_2up_preview(img, page_size, orientation) if layout_pages == 2 else create_image_4up_preview(img, page_size, orientation)
                else:
                    base_img = create_image_single_page_preview(img, page_size, orientation)
                    rows, cols = get_grid_for_layout(layout_pages)
                    preview_img = create_generic_grid_preview(base_img, rows, cols)
                
                # Apply color mode conversion
                preview_img = apply_color_mode(preview_img, color_mode)
                
                preview_img.save(preview_path, "PNG", quality=95)
                return str(preview_path)
                
            except Exception as e:
                print(f"Error processing image: {e}")
                return create_preview_image(local_path, file_type, page_range, page_size, orientation, color_mode)
        
        else:
            # Fallback to basic preview
            return create_preview_image(local_path, file_type, page_range, page_size, orientation, color_mode)
            
    except Exception as e:
        print(f"Error creating layout preview: {e}")
        return create_preview_image(local_path, file_type, page_range, page_size, orientation, color_mode)
    finally:
        if is_temp and os.path.exists(local_path):
            try: os.remove(local_path)
            except: pass

def create_preview_image(file_path, file_type, page_range="1", page_size="A4", orientation="Portrait", color_mode="Color"):
    """
    Create a preview image for the document.
    Handles both local paths and Cloudinary URLs.
    """
    # ROOT FIX: Ensure local path
    local_path, is_temp = ensure_local_path(file_path)
    try:
        # Step 0: Normalize document to PDF for high-fidelity previewing
        DOCUMENT_TYPES = ['docx', 'doc', 'pptx', 'xlsx', 'odt', 'ods', 'odp']

        if file_type.lower() in DOCUMENT_TYPES:
            normalized_path = normalize_document_for_preview(local_path, file_type)
            if normalized_path != local_path:
                local_path = normalized_path
                file_type = 'pdf'
                
        # Create preview directory
        if is_url(file_path):
            preview_dir = UPLOAD_FOLDER / "previews"
        else:
            preview_dir = Path(local_path).parent / "previews"
        preview_dir.mkdir(exist_ok=True)
        
        # Create unique filename with orientation and page size info for better caching
        orientation_suffix = "L" if orientation == "Landscape" else "P"
        page_size_suffix = page_size.replace(" ", "")
        preview_filename = f"preview_{orientation_suffix}_{page_size_suffix}_{uuid.uuid4().hex[:8]}.png"
        preview_path = preview_dir / preview_filename
        
        if file_type == 'pdf':
            # For PDF, try to extract actual page content with customization
            try:
                import fitz  # PyMuPDF for better PDF handling
                
                # Open PDF document
                pdf_document = fitz.open(local_path)
                
                # Parse page range
                page_numbers = parse_page_range(page_range, len(pdf_document))
                
                if not page_numbers:
                    page_numbers = [0]  # Default to first page
                
                # Get the first page from the range
                page_index = page_numbers[0] - 1  # Convert to 0-based index
                
                if 0 <= page_index < len(pdf_document):
                    page = pdf_document[page_index]
                    
                    # Apply page size and orientation transformations
                    page_rect = page.rect
                    
                    # Calculate zoom based on page size
                    zoom_factor = get_zoom_for_page_size(page_size, page_rect)
                    
                    # Apply orientation transformation
                    print(f"PDF Orientation Debug: Requested={orientation}, Page size={page_rect.width}x{page_rect.height}")
                    
                    if orientation == "Landscape":
                        # Force landscape orientation - rotate if page is taller than wide
                        if page_rect.width < page_rect.height:
                            # Page is portrait, rotate to landscape
                            print("Rotating PDF page to landscape")
                            mat = fitz.Matrix(0, 1, -1, 0, page_rect.height, 0)
                            mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
                        else:
                            # Page is already landscape
                            print("PDF page already in landscape")
                            mat = fitz.Matrix(zoom_factor, zoom_factor)
                    else:
                        # Portrait orientation - rotate if page is wider than tall
                        if page_rect.width > page_rect.height:
                            # Page is landscape, rotate to portrait
                            print("Rotating PDF page to portrait")
                            mat = fitz.Matrix(0, -1, 1, 0, 0, page_rect.width)
                            mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
                        else:
                            # Page is already portrait
                            print("PDF page already in portrait")
                            mat = fitz.Matrix(zoom_factor, zoom_factor)
                    
                    # Render page to image with customization
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    
                    # Convert to PIL Image
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Resize to preview size while maintaining aspect ratio
                    img.thumbnail((600, 800), Image.Resampling.LANCZOS)
                    
                    # Apply color mode conversion
                    img = apply_color_mode(img, color_mode)
                    
                    img.save(preview_path, "PNG", quality=95)
                    
                    pdf_document.close()
                else:
                    raise Exception(f"Page {page_index + 1} not found in PDF")
                    
            except ImportError:
                # Fallback to PyPDF2 if PyMuPDF not available
                try:
                    with open(local_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        if pdf_reader.pages:
                            # Create a simple preview with PDF info
                            img = Image.new('RGB', (400, 600), color='white')
                            draw = ImageDraw.Draw(img)
                            try:
                                font = ImageFont.truetype("arial.ttf", 16)
                            except:
                                font = ImageFont.load_default()
                            
                            draw.text((20, 20), f"PDF Document", fill='black', font=font)
                            draw.text((20, 50), f"File: {Path(local_path).name}", fill='gray', font=font)
                            draw.text((20, 80), f"Pages: {len(pdf_reader.pages)}", fill='gray', font=font)
                            draw.text((20, 110), f"Size: {page_size}", fill='gray', font=font)
                            draw.text((20, 140), f"Orientation: {orientation}", fill='gray', font=font)
                            draw.text((20, 170), f"Range: {page_range}", fill='gray', font=font)
                            
                            # Add a simple PDF icon representation
                            draw.rectangle([50, 250, 350, 550], outline='black', width=2)
                            draw.text((150, 400), "PDF", fill='black', font=font)
                            
                            img.save(preview_path)
                        else:
                            raise Exception("No pages in PDF")
                except Exception as e:
                    # Final fallback
                    img = Image.new('RGB', (400, 600), color='white')
                    draw = ImageDraw.Draw(img)
                    try:
                        font = ImageFont.truetype("arial.ttf", 20)
                    except:
                        font = ImageFont.load_default()
                    
                    draw.text((50, 50), f"PDF Preview", fill='black', font=font)
                    draw.text((50, 80), f"File: {Path(local_path).name}", fill='gray', font=font)
                    draw.text((50, 110), f"Error: {str(e)}", fill='red', font=font)
                    
                    img.save(preview_path)
        
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
            # For images, resize and save as preview with customization
            img = Image.open(local_path)
            
            # Convert to RGB if necessary (for JPEG compatibility)
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            
            # Apply orientation transformation
            print(f"Image Orientation Debug: Requested={orientation}, Image size={img.width}x{img.height}")
            
            if orientation == "Landscape":
                # Force landscape orientation - rotate if image is taller than wide
                if img.width < img.height:
                    print("Rotating image to landscape")
                    img = img.rotate(90, expand=True)
                else:
                    print("Image already in landscape")
            else:
                # Portrait orientation - rotate if image is wider than tall
                if img.width > img.height:
                    print("Rotating image to portrait")
                    img = img.rotate(-90, expand=True)
                else:
                    print("Image already in portrait")
            
            # Apply page size scaling
            target_size = get_image_size_for_page_size(page_size)
            img.thumbnail(target_size, Image.Resampling.LANCZOS)
            
            # Resize to preview size while maintaining aspect ratio
            img.thumbnail((600, 800), Image.Resampling.LANCZOS)
            
            # Apply color mode conversion
            img = apply_color_mode(img, color_mode)
            
            img.save(preview_path, "PNG", quality=95)
        
        elif file_type in ['docx', 'doc']:
            # For Word documents, try to extract text content
            try:
                doc = Document(local_path)
                
                # Create preview image with higher resolution
                img = Image.new('RGB', (600, 800), color='white')
                draw = ImageDraw.Draw(img)
                try:
                    font = ImageFont.truetype("arial.ttf", 14)
                    title_font = ImageFont.truetype("arial.ttf", 18)
                except:
                    try:
                        font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 14)
                        title_font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 18)
                    except:
                        font = ImageFont.load_default()
                        title_font = ImageFont.load_default()
                
                # Add title with customization info
                draw.text((30, 30), f"Word Document Preview", fill='black', font=title_font)
                draw.text((30, 60), f"File: {Path(local_path).name}", fill='gray', font=font)
                draw.text((30, 85), f"Page Size: {page_size}", fill='blue', font=font)
                draw.text((30, 110), f"Orientation: {orientation}", fill='blue', font=font)
                if page_range and page_range.strip():
                    draw.text((30, 135), f"Page Range: {page_range}", fill='blue', font=font)
                
                # Extract and display first few paragraphs
                y_pos = 170
                max_lines = 20
                line_count = 0
                
                for paragraph in doc.paragraphs:
                    if line_count >= max_lines:
                        break
                    
                    text = paragraph.text.strip()
                    if text:
                        # Wrap text to fit in preview
                        words = text.split()
                        lines = []
                        current_line = ""
                        
                        for word in words:
                            test_line = current_line + (" " if current_line else "") + word
                            bbox = draw.textbbox((0, 0), test_line, font=font)
                            if bbox[2] - bbox[0] > 540:  # Width limit for higher resolution
                                if current_line:
                                    lines.append(current_line)
                                    current_line = word
                                else:
                                    lines.append(word)
                            else:
                                current_line = test_line
                        
                        if current_line:
                            lines.append(current_line)
                        
                        # Draw lines
                        for line in lines:
                            if y_pos > 750:  # Height limit for higher resolution
                                break
                            draw.text((30, y_pos), line, fill='black', font=font)
                            y_pos += 25
                            line_count += 1
                            
                            if line_count >= max_lines:
                                break
                
                if line_count >= max_lines:
                    draw.text((30, y_pos), "...", fill='gray', font=font)
                
                # Apply color mode conversion
                img = apply_color_mode(img, color_mode)
                
                img.save(preview_path, "PNG", quality=95)
                
            except Exception as e:
                # Fallback for Word documents
                img = Image.new('RGB', (400, 600), color='lightblue')
                draw = ImageDraw.Draw(img)
                try:
                    font = ImageFont.truetype("arial.ttf", 20)
                except:
                    font = ImageFont.load_default()
                
                draw.text((50, 50), f"Word Document Preview", fill='black', font=font)
                draw.text((50, 80), f"File: {Path(local_path).name}", fill='gray', font=font)
                draw.text((50, 110), f"Type: {file_type.upper()}", fill='gray', font=font)
                draw.text((50, 140), f"Error: {str(e)}", fill='red', font=font)
                
                img.save(preview_path)
        
        else:
            # For other files, create a generic placeholder
            img = Image.new('RGB', (400, 600), color='lightgray')
            draw = ImageDraw.Draw(img)
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
            
            draw.text((50, 50), f"Document Preview", fill='black', font=font)
            draw.text((50, 80), f"File: {Path(local_path).name}", fill='gray', font=font)
            draw.text((50, 110), f"Type: {file_type.upper()}", fill='gray', font=font)
            
            img.save(preview_path)
        
        return str(preview_path)
    
    except Exception as e:
        print(f"Error creating preview: {e}")
        return None
    finally:
        if is_temp and os.path.exists(local_path):
            try: os.remove(local_path)
            except: pass

# Layout preview helper functions

def create_single_page_preview(page, page_size, orientation):
    """Create single page preview"""
    try:
        import fitz
        page_rect = page.rect
        zoom_factor = get_zoom_for_page_size(page_size, page_rect)
        
        # Apply orientation
        if orientation == "Landscape":
            if page_rect.width < page_rect.height:
                mat = fitz.Matrix(0, 1, -1, 0, page_rect.height, 0)
                mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
            else:
                mat = fitz.Matrix(zoom_factor, zoom_factor)
        else:
            if page_rect.width > page_rect.height:
                mat = fitz.Matrix(0, -1, 1, 0, 0, page_rect.width)
                mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
            else:
                mat = fitz.Matrix(zoom_factor, zoom_factor)
        
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        img.thumbnail((600, 800), Image.Resampling.LANCZOS)
        return img
    except Exception as e:
        print(f"Error creating single page preview: {e}")
        return create_fallback_preview()

def create_2up_preview(page, page_size, orientation):
    """Create 2-up layout preview (two copies side by side)"""
    try:
        import fitz
        page_rect = page.rect
        zoom_factor = get_zoom_for_page_size(page_size, page_rect) * 0.7  # Smaller for 2-up
        
        # Apply orientation
        if orientation == "Landscape":
            if page_rect.width < page_rect.height:
                mat = fitz.Matrix(0, 1, -1, 0, page_rect.height, 0)
                mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
            else:
                mat = fitz.Matrix(zoom_factor, zoom_factor)
        else:
            if page_rect.width > page_rect.height:
                mat = fitz.Matrix(0, -1, 1, 0, 0, page_rect.width)
                mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
            else:
                mat = fitz.Matrix(zoom_factor, zoom_factor)
        
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_data = pix.tobytes("png")
        single_page = Image.open(io.BytesIO(img_data))
        single_page.thumbnail((300, 400), Image.Resampling.LANCZOS)
        
        # Create 2-up layout
        width, height = single_page.size
        layout_img = Image.new('RGB', (width * 2 + 20, height), 'white')
        
        # Place two copies side by side
        layout_img.paste(single_page, (0, 0))
        layout_img.paste(single_page, (width + 20, 0))
        
        return layout_img
    except Exception as e:
        print(f"Error creating 2-up preview: {e}")
        return create_fallback_preview()

def create_4up_preview(page, page_size, orientation):
    """Create 4-up layout preview (four copies in 2x2 grid)"""
    try:
        import fitz
        page_rect = page.rect
        zoom_factor = get_zoom_for_page_size(page_size, page_rect) * 0.5  # Smaller for 4-up
        
        # Apply orientation
        if orientation == "Landscape":
            if page_rect.width < page_rect.height:
                mat = fitz.Matrix(0, 1, -1, 0, page_rect.height, 0)
                mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
            else:
                mat = fitz.Matrix(zoom_factor, zoom_factor)
        else:
            if page_rect.width > page_rect.height:
                mat = fitz.Matrix(0, -1, 1, 0, 0, page_rect.width)
                mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
            else:
                mat = fitz.Matrix(zoom_factor, zoom_factor)
        
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_data = pix.tobytes("png")
        single_page = Image.open(io.BytesIO(img_data))
        single_page.thumbnail((250, 300), Image.Resampling.LANCZOS)
        
        # Create 4-up layout (2x2 grid)
        width, height = single_page.size
        layout_img = Image.new('RGB', (width * 2 + 20, height * 2 + 20), 'white')
        
        # Place four copies in 2x2 grid
        layout_img.paste(single_page, (0, 0))
        layout_img.paste(single_page, (width + 20, 0))
        layout_img.paste(single_page, (0, height + 20))
        layout_img.paste(single_page, (width + 20, height + 20))
        
        return layout_img
    except Exception as e:
        print(f"Error creating 4-up preview: {e}")
        return create_fallback_preview()

def create_word_single_page_preview(text_content, page_size, orientation):
    """Create single page Word document preview"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import ImageReader
        
        # Choose page size via helper
        page_size_tuple = get_page_size_dimensions(page_size)
        if orientation == "Landscape":
            page_size_tuple = (page_size_tuple[1], page_size_tuple[0])
        
        # Create a temporary PDF
        temp_pdf = io.BytesIO()
        c = canvas.Canvas(temp_pdf, pagesize=page_size_tuple)
        width, height = page_size_tuple
        
        # Add text content
        c.setFont("Helvetica", 12)
        y_position = height - 50
        
        lines = text_content.split('\n')
        for line in lines[:30]:  # Limit to 30 lines
            if y_position < 50:
                break
            c.drawString(50, y_position, line[:80])  # Limit line length
            y_position -= 20
        
        c.save()
        temp_pdf.seek(0)
        
        # Convert to image
        import fitz
        doc = fitz.open(stream=temp_pdf.read(), filetype="pdf")
        page = doc[0]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        img.thumbnail((600, 800), Image.Resampling.LANCZOS)
        doc.close()
        
        return img
    except Exception as e:
        print(f"Error creating Word single page preview: {e}")
        return create_fallback_preview()

def create_word_2up_preview(text_content, page_size, orientation):
    """Create 2-up Word document preview"""
    single_page = create_word_single_page_preview(text_content, page_size, orientation)
    single_page.thumbnail((300, 400), Image.Resampling.LANCZOS)
    
    # Create 2-up layout
    width, height = single_page.size
    layout_img = Image.new('RGB', (width * 2 + 20, height), 'white')
    layout_img.paste(single_page, (0, 0))
    layout_img.paste(single_page, (width + 20, 0))
    
    return layout_img

def create_word_4up_preview(text_content, page_size, orientation):
    """Create 4-up Word document preview"""
    single_page = create_word_single_page_preview(text_content, page_size, orientation)
    single_page.thumbnail((250, 300), Image.Resampling.LANCZOS)
    
    # Create 4-up layout
    width, height = single_page.size
    layout_img = Image.new('RGB', (width * 2 + 20, height * 2 + 20), 'white')
    layout_img.paste(single_page, (0, 0))
    layout_img.paste(single_page, (width + 20, 0))
    layout_img.paste(single_page, (0, height + 20))
    layout_img.paste(single_page, (width + 20, height + 20))
    
    return layout_img

def create_image_single_page_preview(img, page_size, orientation):
    """Create single page image preview"""
    try:
        # Apply orientation
        if orientation == "Landscape" and img.width < img.height:
            img = img.rotate(90, expand=True)
        elif orientation == "Portrait" and img.width > img.height:
            img = img.rotate(90, expand=True)
        
        # Scale to appropriate size
        img.thumbnail((600, 800), Image.Resampling.LANCZOS)
        return img
    except Exception as e:
        print(f"Error creating image single page preview: {e}")
        return create_fallback_preview()

def create_image_2up_preview(img, page_size, orientation):
    """Create 2-up image preview"""
    single_page = create_image_single_page_preview(img, page_size, orientation)
    single_page.thumbnail((300, 400), Image.Resampling.LANCZOS)
    
    # Create 2-up layout
    width, height = single_page.size
    layout_img = Image.new('RGB', (width * 2 + 20, height), 'white')
    layout_img.paste(single_page, (0, 0))
    layout_img.paste(single_page, (width + 20, 0))
    
    return layout_img

def create_image_4up_preview(img, page_size, orientation):
    """Create 4-up image preview"""
    single_page = create_image_single_page_preview(img, page_size, orientation)
    single_page.thumbnail((250, 300), Image.Resampling.LANCZOS)
    
    # Create 4-up layout
    width, height = single_page.size
    layout_img = Image.new('RGB', (width * 2 + 20, height * 2 + 20), 'white')
    layout_img.paste(single_page, (0, 0))
    layout_img.paste(single_page, (width + 20, 0))
    layout_img.paste(single_page, (0, height + 20))
    layout_img.paste(single_page, (width + 20, height + 20))
    
    return layout_img

def create_fallback_preview():
    """Create a fallback preview when layout creation fails"""
    img = Image.new('RGB', (600, 800), 'white')
    return img

def get_grid_for_layout(layout_pages):
    """Return rows, cols for a given layout_pages following common printer driver defaults."""
    mapping = {
        2: (1, 2),
        4: (2, 2),
        6: (3, 2),
        8: (4, 2),
        9: (3, 3),
        16: (4, 4)
    }
    if layout_pages in mapping:
        return mapping[layout_pages]
    # Fallback: nearest square
    import math
    cols = math.ceil(math.sqrt(layout_pages))
    rows = math.ceil(layout_pages / cols)
    return rows, cols

def create_generic_grid_preview(single_page_img, rows, cols, gap=20, bg_color='white'):
    """Tile a single PIL image into a rows x cols grid for preview."""
    # Compute cell size based on a target preview area
    target_width = min(800, single_page_img.width * cols)
    target_height = min(800, single_page_img.height * rows)
    cell_width = int((target_width - gap * (cols - 1)) / cols)
    cell_height = int((target_height - gap * (rows - 1)) / rows)
    # Preserve aspect ratio of the single page
    aspect = single_page_img.width / single_page_img.height
    if cell_width / cell_height > aspect:
        cell_width = int(cell_height * aspect)
    else:
        cell_height = int(cell_width / aspect)
    layout_img = Image.new('RGB', (cell_width * cols + gap * (cols - 1), cell_height * rows + gap * (rows - 1)), bg_color)
    resized = single_page_img.copy()
    resized.thumbnail((cell_width, cell_height), Image.Resampling.LANCZOS)
    for r in range(rows):
        for c in range(cols):
            x = c * (cell_width + gap)
            y = r * (cell_height + gap)
            layout_img.paste(resized, (x, y))
    return layout_img

# ── Phase 3A: Direct PIL preview for images (no PDF intermediary) ────────
# This function bypasses the ENTIRE PDF pipeline for image previews:
#   ✗ No ReportLab canvas
#   ✗ No intermediate PDF file
#   ✗ No fitz.open() / get_pixmap()
#   ✗ No PDF rasterization
#
# Instead:  Image → PIL → EXIF → orientation → color → canvas → JPEG
#
# ONLY used when:  file_type ∈ {jpg, jpeg, png, webp, bmp, tiff} AND preview_mode=True
# NEVER used for:  print path, PDFs, DOCX, multi-page documents
#
# Rollback: Delete this function + the call site in app.py render endpoint.
#           generate_final_print_pdf() continues to work unchanged.

_IMAGE_PREVIEW_TYPES = {'jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff'}

def generate_image_preview_direct(
    file_path,
    file_type,
    output_dir,
    page_size="A4",
    orientation="Portrait",
    layout_pages=1,
    color_mode="Color",
    max_preview_width=800,
    jpeg_quality=60
):
    """
    Generate preview JPEG(s) directly from an image file using only PIL.
    Completely bypasses ReportLab + fitz for maximum speed.

    Args:
        file_path: Path to the source image file
        file_type: File extension (must be in _IMAGE_PREVIEW_TYPES)
        output_dir: Directory to write preview JPEG files into
        page_size: Target page size for layout composition (e.g. "A4")
        orientation: "Portrait" or "Landscape"
        layout_pages: Pages per sheet (1, 2, 4, 6, 8, 9, 16)
        color_mode: "Color" or "Black & White"
        max_preview_width: Maximum width of output preview JPEG
        jpeg_quality: JPEG compression quality (0-100)

    Returns:
        list[str]: List of absolute paths to generated preview JPEG files

    Raises:
        ValueError: If file_type is not an image type
        FileNotFoundError: If file_path does not exist
    """
    import math
    from PIL import ImageOps

    if file_type not in _IMAGE_PREVIEW_TYPES:
        raise ValueError(f"generate_image_preview_direct only handles images, got: {file_type}")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Source image not found: {file_path}")

    t_start = time.time()

    # ── 1. Load image + EXIF correction ──────────────────────────────
    img = Image.open(file_path)
    img = ImageOps.exif_transpose(img)
    if img.mode != 'RGB':
        img = img.convert('RGB')

    # Thumbnail to preview-adequate size (saves memory in composition)
    img.thumbnail((1600, 1600), Image.Resampling.LANCZOS)

    # ── 2. Apply color mode ──────────────────────────────────────────
    img = apply_color_mode(img, color_mode)

    # ── 3. Apply orientation ─────────────────────────────────────────
    img_is_portrait = img.height > img.width
    if orientation == "Landscape" and img_is_portrait:
        img = img.rotate(90, expand=True)
    elif orientation == "Portrait" and not img_is_portrait:
        img = img.rotate(90, expand=True)

    # ── 4. Get target page dimensions (in pixels for preview) ────────
    # Use 72 DPI equivalent — page_size returns points (1 point = 1 px at 72 DPI)
    page_w_pt, page_h_pt = get_page_size_dimensions(page_size)

    # For 2-per-sheet, output is landscape
    if layout_pages == 2:
        if page_w_pt < page_h_pt:
            page_w_pt, page_h_pt = page_h_pt, page_w_pt

    # Scale to preview size (target ~800px max dimension)
    scale = max_preview_width / max(page_w_pt, page_h_pt)
    canvas_w = int(page_w_pt * scale)
    canvas_h = int(page_h_pt * scale)

    # ── 5. Layout composition ────────────────────────────────────────
    rows, cols = get_grid_for_layout(layout_pages)
    margin = int(18 * scale)  # Match ReportLab margin (18 points)

    grid_w = canvas_w - 2 * margin
    grid_h = canvas_h - 2 * margin

    if layout_pages > 1:
        cell_w = grid_w // cols
        cell_h = grid_h // rows
    else:
        cell_w = grid_w
        cell_h = grid_h

    # Images always have 1 logical page — for N-up layout, place the
    # image in slot 0 only; remaining slots stay blank (white).
    # This matches generate_final_print_pdf's None-padding semantics.
    num_sheets = 1  # Single image = always 1 sheet

    # Create canvas (white background)
    canvas_img = Image.new('RGB', (canvas_w, canvas_h), (255, 255, 255))

    if layout_pages == 1:
        # ── Single page per sheet ────────────────────────────────────
        img_w, img_h = img.size
        img_aspect = img_w / img_h

        if cell_w / cell_h > img_aspect:
            draw_h = cell_h
            draw_w = int(cell_h * img_aspect)
        else:
            draw_w = cell_w
            draw_h = int(cell_w / img_aspect)

        resized = img.resize((draw_w, draw_h), Image.Resampling.LANCZOS)

        # Center on canvas
        x = margin + (cell_w - draw_w) // 2
        y = margin + (cell_h - draw_h) // 2

        canvas_img.paste(resized, (x, y))
    else:
        # ── Multi-page layout (page-aware, Xerox N-up semantics) ─────
        # 1 image = 1 logical document page.  Build a slot list padded
        # with None so unfilled slots remain blank (white background).
        # This mirrors generate_final_print_pdf's sheet_page_nums logic.
        slot_pages = [img] + [None] * (layout_pages - 1)

        img_w, img_h = img.size
        img_aspect = img_w / img_h

        if cell_w / cell_h > img_aspect:
            draw_h = cell_h
            draw_w = int(cell_h * img_aspect)
        else:
            draw_w = cell_w
            draw_h = int(cell_w / img_aspect)

        resized = img.resize((draw_w, draw_h), Image.Resampling.LANCZOS)

        for idx, page in enumerate(slot_pages):
            if page is None:
                continue  # Blank slot — white background shows through

            r = idx // cols
            col_idx = idx % cols

            x = margin + col_idx * cell_w + (cell_w - draw_w) // 2
            y = margin + r * cell_h + (cell_h - draw_h) // 2

            canvas_img.paste(resized, (x, y))

    # ── 6. Downscale if needed ───────────────────────────────────────
    if canvas_img.width > max_preview_width:
        ratio = max_preview_width / canvas_img.width
        canvas_img = canvas_img.resize(
            (max_preview_width, int(canvas_img.height * ratio)),
            Image.Resampling.LANCZOS
        )

    # ── 7. Save preview JPEG directly (NO intermediate PDF) ──────────
    os.makedirs(output_dir, exist_ok=True)
    preview_filename = f"preview_{uuid.uuid4().hex[:8]}_sheet_1.jpg"
    preview_path = os.path.join(output_dir, preview_filename)

    canvas_img.save(preview_path, "JPEG", quality=jpeg_quality, optimize=True)

    t_elapsed = time.time() - t_start
    logger.info(f"IMAGE PREVIEW DIRECT: {file_type} rendered in {t_elapsed:.3f}s "
                f"(layout={layout_pages}, orient={orientation}, color={color_mode}, "
                f"size={canvas_img.width}x{canvas_img.height})")

    return [preview_path]


def generate_final_print_pdf(file_path, file_type, page_size="A4", orientation="Portrait", layout_pages=1, color_mode="Color", page_range="", preview_mode=False, color_pages=None):
    """
    Generate final print-ready PDF that matches preview output EXACTLY.
    Handles both local paths and Cloudinary URLs.
    
    Phase 3D: Per-page color rendering.
    If color_pages is provided (list of 1-indexed page numbers designated as color),
    pages IN the list are rendered in Color, pages NOT in the list are rendered
    in Black & White. If color_pages is None, the global color_mode is applied
    uniformly to all pages (preserving existing behavior).
    """
    # ROOT FIX: Ensure local path
    src_local_path, is_temp = ensure_local_path(file_path)
    
    try:
        import math
        
        # Phase 1E-A: Architectural isolation -- detect image-only preview requests.
        # This flag enables future phases to optimize the image preview path
        # independently without touching the generic PDF/document pipeline.
        _IMAGE_PREVIEW_TYPES = {'jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff'}
        _is_image_preview = preview_mode and file_type in _IMAGE_PREVIEW_TYPES

        # Phase 3D: Parse color_pages into a fast-lookup set.
        # If color_pages is provided and non-empty, per-page rendering activates.
        # Otherwise, the global color_mode path is used (zero overhead).
        _color_pages_set = None
        if color_pages:
            try:
                import json as _json
                if isinstance(color_pages, str):
                    _parsed = _json.loads(color_pages)
                elif isinstance(color_pages, (list, set, tuple)):
                    _parsed = list(color_pages)
                else:
                    _parsed = None
                if _parsed and isinstance(_parsed, list):
                    _color_pages_set = set(int(p) for p in _parsed
                                           if isinstance(p, (int, float)) and p > 0)
                    if not _color_pages_set:
                        _color_pages_set = None
            except Exception:
                _color_pages_set = None

        _is_mixed_render = _color_pages_set is not None
        if _is_mixed_render:
            logger.info(f"PHASE 3D: Mixed rendering activated -- "
                        f"{len(_color_pages_set)} color pages, "
                        f"remaining pages will be B&W")

        # Step 0: Normalize non-PDF files for high-fidelity conversion
        pdf_path = src_local_path
        pdf_document = None
        if _is_image_preview:
            # ── Phase 1E-C: Complete normalization bypass for image previews ──
            # Images are rendered directly via PIL (Phase 1E-B) — no need for
            # PDF conversion or fitz.  Skip normalize_document_for_preview()
            # and fitz.open() entirely.  Images always have exactly 1 page.
            # Print path and PDF/doc preview path never reach here.
            logger.debug(f"IMAGE PREVIEW PATH: Bypassing normalization for {file_type}")
            total_pages = 1
            selected_page_nums = [1]
        else:
            # ── Generic document/PDF + print path (unchanged) ──
            if file_type != 'pdf':
                pdf_path = normalize_document_for_preview(src_local_path, file_type)
                if pdf_path == src_local_path:
                    print(f"Warning: Could not normalize {file_type} to PDF for final print")
                else:
                    file_type = 'pdf'

            # Step 1: Open PDF document
            try:
                import fitz
                pdf_document = fitz.open(pdf_path)
                total_pages = len(pdf_document)
            except Exception as e:
                print(f"Error opening PDF: {e}")
                return file_path

            # Parse page range to get selected pages
            if page_range and page_range.strip():
                try:
                    selected_page_nums = parse_page_range(page_range.strip(), total_pages)
                except Exception as e:
                    print(f"Error parsing page range '{page_range}': {e}, using all pages")
                    selected_page_nums = list(range(1, total_pages + 1))
            else:
                selected_page_nums = list(range(1, total_pages + 1))

            # CRITICAL SAFETY CHECK: Reject job if page range results in 0 pages
            if not selected_page_nums:
                pdf_document.close()
                error_msg = f"Page range '{page_range}' is invalid or exceeds document page count ({total_pages} pages)"
                print(f"ERROR: {error_msg}")
                raise ValueError(error_msg)
        
        # Get grid dimensions for layout
        rows, cols = get_grid_for_layout(layout_pages)
        
        # Calculate number of sheets needed
        num_sheets = math.ceil(len(selected_page_nums) / layout_pages) if layout_pages > 1 else len(selected_page_nums)
        
        # Get page dimensions
        page_w, page_h = get_page_size_dimensions(page_size)
        
        
        # 2-per-sheet → output page landscape hona chahiye (printer driver behavior)
        if layout_pages == 2:
            if page_w < page_h:  # portrait hai → landscape karo
                page_w, page_h = page_h, page_w
        
        logger.info(f"DEBUG OUTPUT: page_w={page_w:.1f} page_h={page_h:.1f} layout_pages={layout_pages}")
        # Create output PDF
        output_pdf = str(Path(file_path).with_name(f"final_print_{uuid.uuid4().hex[:8]}.pdf"))
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.utils import ImageReader
        
        c = rl_canvas.Canvas(output_pdf, pagesize=(page_w, page_h))
        
        # Compute grid cell size with margins (matching preview layout)
        margin = 18  # 0.25 inch
        grid_w = page_w - 2 * margin
        grid_h = page_h - 2 * margin
        
        if layout_pages > 1:
            cell_w = grid_w / cols
            cell_h = grid_h / rows
        else:
            cell_w = grid_w
            cell_h = grid_h
        
        # Process each sheet
        for sheet_idx in range(num_sheets):
            # Get pages for this sheet
            start_idx = sheet_idx * layout_pages
            end_idx = min(start_idx + layout_pages, len(selected_page_nums))
            sheet_page_nums = selected_page_nums[start_idx:end_idx]
            
            # Pad with white placeholders if needed (matching preview behavior)
            while len(sheet_page_nums) < layout_pages:
                sheet_page_nums.append(None)  # None = white placeholder
            
            # Process each page in this sheet
            page_images = []
            for page_num in sheet_page_nums:
                if page_num is None:
                    # White placeholder
                    page_images.append(None)
                    continue
                
                try:
                    if _is_image_preview:
                        # ── Phase 1E-B: Direct PIL rendering for image previews ──
                        # Bypasses fitz PDF rasterization entirely.
                        # Loads the original image directly with PIL, applies EXIF
                        # correction, color mode, and orientation — then feeds the
                        # result into the SAME layout composition code below.
                        # Print path never reaches here (_is_image_preview=False).
                        from PIL import ImageOps
                        img = Image.open(src_local_path)
                        img = ImageOps.exif_transpose(img)
                        if img.mode != 'RGB':
                            img = img.convert('RGB')

                        # Thumbnail to preview-adequate size (saves memory/CPU
                        # in downstream layout composition).
                        img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)

                        # Phase 3D: Per-page color mode for image previews
                        _eff_cm = color_mode
                        if _is_mixed_render and page_num is not None:
                            _eff_cm = 'Color' if page_num in _color_pages_set else 'Black & White'
                        img = apply_color_mode(img, _eff_cm)

                        # Orientation: match existing logic using image dimensions
                        img_is_portrait = img.height > img.width

                        if orientation == "Landscape" and img_is_portrait:
                            img = img.rotate(90, expand=True)
                        elif orientation == "Portrait" and not img_is_portrait:
                            img = img.rotate(90, expand=True)

                        page_images.append(img)
                    else:
                        # ── Generic fitz path (PDFs, documents, print) ── unchanged
                        # Get page (1-based to 0-based index)
                        page_idx = page_num - 1
                        if 0 <= page_idx < len(pdf_document):
                            page = pdf_document[page_idx]
                            page_rect = page.rect
                            
                            # Phase 1D-B: Reduce rasterization cost for previews.
                            # Print path (preview_mode=False): 2.0x — full fidelity, unchanged.
                            # Preview path (preview_mode=True): 0.5x — sufficient for 800px output.
                            zoom_factor = 0.5 if preview_mode else 2.0
                            mat = fitz.Matrix(zoom_factor, zoom_factor)
                            
                            # Render page to image
                            pix = page.get_pixmap(matrix=mat, alpha=False)
                            img_data = pix.tobytes("png")
                            img = Image.open(io.BytesIO(img_data))
                            # Phase 3D: Per-page color mode for PDF/document rendering
                            _eff_cm = color_mode
                            if _is_mixed_render:
                                _eff_cm = 'Color' if page_num in _color_pages_set else 'Black & White'
                            img = apply_color_mode(img, _eff_cm)

                            img_w, img_h = img.size
                            # Use original PDF page orientation (more reliable)
                            page_is_portrait = page_rect.height > page_rect.width

                            if orientation == "Landscape" and page_is_portrait:
                                img = img.rotate(90, expand=True)

                            elif orientation == "Portrait" and not page_is_portrait:
                                img = img.rotate(90, expand=True)           

                            page_images.append(img)
                        else:
                            page_images.append(None)  # Out of bounds, use placeholder
                except Exception as e:
                    print(f"Error processing page {page_num}: {e}")
                    page_images.append(None)  # Error, use placeholder
            
            # Draw this sheet
            if layout_pages == 1:
                # Single page per sheet
                if page_images[0]:
                    # Calculate size to fit page while maintaining aspect ratio
                    img_w, img_h = page_images[0].size
                    img_aspect = img_w / img_h

                    # Save AFTER rotation
                    tmp_img_path = str(Path(file_path).with_name(f"print_img_{uuid.uuid4().hex[:8]}.png"))
                    page_images[0].save(tmp_img_path, "PNG")
                    img_reader = ImageReader(tmp_img_path)

                    if cell_w / cell_h > img_aspect:
                        draw_w = cell_h * img_aspect
                        draw_h = cell_h
                    else:
                        draw_w = cell_w
                        draw_h = cell_w / img_aspect
                    
                    # Center on page
                    x = margin + (cell_w - draw_w) / 2
                    y = page_h - (margin + cell_h) + (cell_h - draw_h) / 2
                    
                    c.drawImage(img_reader, x, y, width=draw_w, height=draw_h, preserveAspectRatio=True, mask='auto')
                    
                    # Cleanup
                    try:
                        os.remove(tmp_img_path)
                    except Exception:
                        pass
            else:
                # Multiple pages per sheet (layout)
                for idx, page_img in enumerate(page_images):
                    if page_img is None:
                        continue  # Skip placeholders (white background)
                    
                    # Calculate grid position
                    r = idx // cols
                    col = idx % cols
                    
                    # Calculate size to fit cell while maintaining aspect ratio
                    img_w, img_h = page_img.size
                    img_aspect = img_w / img_h

                    # Save AFTER rotation
                    tmp_img_path = str(Path(file_path).with_name(f"print_img_{uuid.uuid4().hex[:8]}.png"))
                    page_img.save(tmp_img_path, "PNG")
                    img_reader = ImageReader(tmp_img_path)

                    if cell_w / cell_h > img_aspect:
                        draw_w = cell_h * img_aspect
                        draw_h = cell_h
                    else:
                        draw_w = cell_w
                        draw_h = cell_w / img_aspect
                    
                    # Calculate position in grid
                    x = margin + col * cell_w + (cell_w - draw_w) / 2
                    y = page_h - (margin + (r + 1) * cell_h) + (cell_h - draw_h) / 2
                    
                    c.drawImage(img_reader, x, y, width=draw_w, height=draw_h, preserveAspectRatio=True, mask='auto')
                    
                    # Cleanup
                    try:
                        os.remove(tmp_img_path)
                    except Exception:
                        pass
            
            # Finish this sheet
            c.showPage()
        
        # Save PDF
        c.save()
        if pdf_document:
            pdf_document.close()
        
        # Cleanup converted PDF if temporary
        if pdf_path != src_local_path and Path(pdf_path).exists():
            try:
                os.remove(pdf_path)
            except Exception:
                pass
        
        return output_pdf
        
    except Exception as e:
        print(f"Error generating final print PDF: {e}")
        import traceback
        traceback.print_exc()
        return file_path
    finally:
        if is_temp and os.path.exists(src_local_path):
            try: os.remove(src_local_path)
            except: pass

def generate_nup_pdf(file_path, file_type, page_size="A4", orientation="Portrait", layout_pages=1, color_mode="Color"):
    """
    LEGACY: Generate a printable single-sheet N-up PDF from the first page/image.
    DEPRECATED: Use generate_final_print_pdf() instead for preview-print matching.
    Returns path to the generated PDF (or original path if layout_pages == 1).
    """
    try:
        if layout_pages == 1:
            return file_path
        rows, cols = get_grid_for_layout(layout_pages)
        # Prepare base image from the first page
        base_img = None
        if file_type == 'pdf':
            try:
                import fitz
                pdf_document = fitz.open(file_path)
                page = pdf_document[0]
                page_rect = page.rect
                zoom_factor = 2.0  # high-res rasterization for print
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom_factor, zoom_factor), alpha=False)
                img_data = pix.tobytes("png")
                base_img = Image.open(io.BytesIO(img_data))
                pdf_document.close()
            except Exception as e:
                print(f"Error rasterizing PDF for N-up: {e}")
                return file_path
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
            base_img = Image.open(file_path)
        elif file_type in ['docx', 'doc']:
            # Reuse word single page preview as base
            doc = Document(file_path) if file_type == 'docx' else None
            text_content = "\n".join([p.text for p in doc.paragraphs]) if doc else ""
            base_img = create_word_single_page_preview(text_content, page_size, orientation)
        if base_img is None:
            return file_path
        # Apply color mode
        base_img = apply_color_mode(base_img, color_mode)
        # Create PDF canvas
        from reportlab.pdfgen import canvas as rl_canvas
        page_w, page_h = get_page_size_dimensions(page_size)
        if orientation == 'Landscape':
            page_w, page_h = page_h, page_w
        output_pdf = str(Path(file_path).with_name(f"nup_{layout_pages}up_{uuid.uuid4().hex[:8]}.pdf"))
        c = rl_canvas.Canvas(output_pdf, pagesize=(page_w, page_h))
        # Compute grid cell size with small margins
        margin = 18  # 0.25 inch
        grid_w = page_w - 2 * margin
        grid_h = page_h - 2 * margin
        cell_w = grid_w / cols
        cell_h = grid_h / rows
        # Maintain aspect ratio of base image
        img_aspect = base_img.width / base_img.height
        if cell_w / cell_h > img_aspect:
            draw_w = cell_h * img_aspect
            draw_h = cell_h
        else:
            draw_w = cell_w
            draw_h = cell_w / img_aspect
        # Convert PIL image to a temporary file for ReportLab
        tmp_img_path = str(Path(file_path).with_name(f"nup_img_{uuid.uuid4().hex[:8]}.png"))
        base_img.save(tmp_img_path, "PNG")
        from reportlab.lib.utils import ImageReader
        img_reader = ImageReader(tmp_img_path)
        # Draw grid
        for r in range(rows):
            for col in range(cols):
                x = margin + col * cell_w + (cell_w - draw_w) / 2
                y = page_h - (margin + (r + 1) * cell_h) + (cell_h - draw_h) / 2
                c.drawImage(img_reader, x, y, width=draw_w, height=draw_h, preserveAspectRatio=True, mask='auto')
        c.showPage()
        c.save()
        # Cleanup temp image
        try:
            os.remove(tmp_img_path)
        except Exception:
            pass
        return output_pdf
    except Exception as e:
        print(f"Error generating N-up PDF: {e}")
        return file_path

def apply_color_mode(img, color_mode):
    """Apply color mode conversion to image"""
    try:
        if color_mode == "Black & White":
            # Convert to grayscale
            if img.mode != 'L':
                img = img.convert('L')
                # Convert back to RGB for consistency
                img = img.convert('RGB')
        # If color_mode is "Color", return original image
        return img
    except Exception as e:
        print(f"Error applying color mode: {e}")
        return img

def get_page_size_dimensions(page_size):
    """
    Get dimensions for page size
    
    Args:
        page_size (str): Page size name
    
    Returns:
        tuple: (width, height) in points
    """
    PAGE_SIZES_PT = {
        'A3': A3,
        'A4': A4,
        'Letter': letter,
        'Legal': legal,
        'Tabloid': (792, 1224),
        'Executive': (522, 756),
        'Envelope Monarch': (279, 540),
        'Envelope #10': (297, 684),
        'B4 (JIS)': (728, 1032),
        'B5 (JIS)': (516, 728),
    }
    return PAGE_SIZES_PT.get(page_size, A4)

def parse_page_range(page_range, total_pages):
    """
    Parse page range string and return list of page numbers.
    
    PAGE RANGE FIX: Enhanced error handling for invalid inputs.
    
    Args:
        page_range (str): Page range like "1-3", "1,3,5", "1-3,7,9-10"
        total_pages (int): Total number of pages in document
    
    Returns:
        list: List of page numbers (1-based)
    
    Raises:
        ValueError: If page range contains invalid format (non-numeric, invalid ranges)
    """
    if not page_range or page_range.strip() == "":
        return list(range(1, total_pages + 1))
    
    pages = []
    parts = page_range.split(',')
    
    for part in parts:
        part = part.strip()
        if not part:
            continue  # Skip empty parts
        
        if '-' in part:
            # Range like "1-3"
            range_parts = part.split('-', 1)
            if len(range_parts) != 2:
                raise ValueError(f"Invalid range format: '{part}'")
            
            try:
                start = int(range_parts[0].strip())
                end = int(range_parts[1].strip())
            except ValueError:
                raise ValueError(f"Range contains non-numeric values: '{part}'")
            
            if start <= 0 or end <= 0:
                raise ValueError(f"Page numbers must be positive: '{part}'")
            
            if start > end:
                raise ValueError(f"Range start must be <= end: '{part}'")
            
            pages.extend(range(start, end + 1))
        else:
            # Single page
            try:
                page_num = int(part)
            except ValueError:
                raise ValueError(f"Invalid page number: '{part}'")
            
            if page_num <= 0:
                raise ValueError(f"Page number must be positive: '{part}'")
            
            pages.append(page_num)
    
    # Filter out invalid page numbers (out of bounds)
    valid_pages = [p for p in pages if 1 <= p <= total_pages]
    return sorted(list(set(valid_pages)))  # Remove duplicates and sort

def get_zoom_for_page_size(page_size, page_rect):
    """
    Calculate zoom factor based on page size setting
    
    Args:
        page_size (str): Target page size (A4, A3, Letter, Legal)
        page_rect: PDF page rectangle
    
    Returns:
        float: Zoom factor
    """
    # Standard page dimensions in points
    page_sizes = {
        'A3': (842, 1191),
        'A4': (595, 842),
        'Letter': (612, 792),
        'Legal': (612, 1008),
        'Tabloid': (792, 1224),
        'Executive': (522, 756),
        'Envelope Monarch': (279, 540),
        'Envelope #10': (297, 684),
        'B4 (JIS)': (728, 1032),
        'B5 (JIS)': (516, 728),
    }
    
    if page_size in page_sizes:
        target_width, target_height = page_sizes[page_size]
        current_width, current_height = page_rect.width, page_rect.height
        
        # Calculate zoom to fit target size
        zoom_x = target_width / current_width
        zoom_y = target_height / current_height
        
        # Use the smaller zoom to ensure it fits
        zoom = min(zoom_x, zoom_y)
        
        # Limit zoom to reasonable range
        return max(0.5, min(zoom, 3.0))
    
    return 1.0  # Default zoom

def get_image_size_for_page_size(page_size):
    """
    Get target image size based on page size setting
    
    Args:
        page_size (str): Target page size (A4, A3, Letter, Legal)
    
    Returns:
        tuple: (width, height) in pixels
    """
    # Standard page dimensions in pixels (at 300 DPI)
    page_sizes = {
        'A3': (3508, 4961),
        'A4': (2480, 3508),
        'Letter': (2550, 3300),
        'Legal': (2550, 4200),
        'Tabloid': (3300, 5100),
        'Executive': (2175, 3150),
        'Envelope Monarch': (1163, 2250),
        'Envelope #10': (1238, 2850),
        'B4 (JIS)': (3035, 4299),
        'B5 (JIS)': (2150, 3035),
    }
    
    if page_size in page_sizes:
        return page_sizes[page_size]
    
    return (2480, 3508)  # Default to A4

def normalize_document_for_preview(file_path, file_type):
    """
    Normalization step for non-PDF documents to ensure they can be previewed.
    Handles both local paths and Cloudinary URLs.
    """
    if file_type == 'pdf':
        return file_path

    # Bypass: images are handled natively by PIL/ReportLab — no LibreOffice needed
    if file_type.lower() in {'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'tiff'}:
        return file_path
        
    try:
        # Create cache directory
        cache_dir = UPLOAD_FOLDER / "preview_cache"
        cache_dir.mkdir(exist_ok=True)
        
        import hashlib
        import shutil

        # ADD THESE 2 LINES:
        soffice_path = shutil.which('soffice') or shutil.which('libreoffice')
        print(f"DEBUG: LibreOffice found at: {soffice_path}")
        
        # ROOT FIX: Handle cache key for URLs vs Local Files
        if is_url(file_path):
            # For URLs, use the URL itself as the cache identity
            cache_key = hashlib.md5(file_path.encode()).hexdigest()
        else:
            # For local files, use path + metadata
            file_stat = os.stat(file_path)
            cache_key = hashlib.md5(f"{file_path}_{file_stat.st_mtime}_{file_stat.st_size}".encode()).hexdigest()
        
        cached_pdf = cache_dir / f"{cache_key}.pdf"
        
        # Return from cache if already converted
        if cached_pdf.exists():
            return str(cached_pdf)
        
        # ROOT FIX: Ensure local path for conversion tool
        local_path, is_temp = ensure_local_path(file_path)
        try:
            print(f"DEBUG: Normalizing {file_type} for preview: {Path(local_path).name}")
            converted_pdf = convert_to_pdf(local_path, file_type)
            
            if converted_pdf != local_path and os.path.exists(converted_pdf):
                import shutil
                shutil.copy2(converted_pdf, str(cached_pdf))
                # Cleanup intermediate
                if str(Path(converted_pdf).parent) != str(cache_dir):
                    try: os.remove(converted_pdf)
                    except: pass
                return str(cached_pdf)
        finally:
            if is_temp and os.path.exists(local_path):
                try: os.remove(local_path)
                except: pass
                
        return file_path
    except Exception as e:
        print(f"Error during document normalization: {e}")
        return file_path

def convert_to_pdf(file_path, file_type):
    """
    Convert non-PDF files to PDF for preview generation.
    Supports: DOCX, DOC, ODT, PPTX, PPT, XLSX, XLS
    
    Args:
        file_path (str): Path to the file
        file_type (str): File extension
    
    Returns:
        str: Path to converted PDF file, or original path if already PDF or conversion fails
    """
    if file_type == 'pdf':
        return file_path
    
    try:
        # Try LibreOffice headless conversion first (industry standard for quality)
        libreoffice_paths = []

        # 1️⃣ First priority: ENV override (production-safe)
        env_path = os.environ.get("LIBREOFFICE_PATH")
        if env_path:
            libreoffice_paths.append(env_path)

        # 2️⃣ OS-specific fallbacks
        if os.name == "nt":  # Windows
            libreoffice_paths.extend([
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                r"C:\LibreOffice\program\soffice.exe"
            ])
        else:
            # Linux / Render
            libreoffice_paths.append("soffice")

        
        libreoffice_exe = None
        for path in libreoffice_paths:
            if not path: continue
            try:
                # Check version as a proxy for existence and availability
                result = subprocess.run([path, '--version'], capture_output=True, timeout=3)
                if result.returncode == 0:
                    libreoffice_exe = path
                    break
            except (FileNotFoundError, subprocess.TimeoutExpired, Exception):
                continue
        
        # 1. Try LibreOffice if available
        if libreoffice_exe:
            print(f"DEBUG: Using LibreOffice {libreoffice_exe} for conversion")
            output_dir = tempfile.mkdtemp()
            try:
                # Convert using headless LibreOffice
                cmd = [
                    libreoffice_exe,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    file_path
                ]
                result = subprocess.run(cmd, capture_output=True, timeout=45)
                if result.returncode == 0:
                    converted_files = list(Path(output_dir).glob("*.pdf"))
                    if converted_files:
                        pdf_path = converted_files[0]
                        final_pdf_path = Path(file_path).parent / f"norm_{uuid.uuid4().hex[:8]}.pdf"
                        import shutil
                        shutil.copy2(str(pdf_path), str(final_pdf_path))
                        return str(final_pdf_path)
            except Exception as e:
                print(f"LibreOffice conversion failed: {e}")
            finally:
                import shutil
                shutil.rmtree(output_dir, ignore_errors=True)
        
        # 2. Try Microsoft Word on Windows if LibreOffice is missing (High Fidelity)
        if os.name == 'nt' and file_type in ['docx', 'doc']:
            print(f"DEBUG: Attempting Microsoft Word conversion for {file_type}")
            try:
                import pythoncom
                from win32com import client
                
                pythoncom.CoInitialize()
                try:
                    word = client.DispatchEx("Word.Application")
                    word.Visible = False
                    
                    # Absolute paths are required for COM
                    abs_input = str(Path(file_path).absolute())
                    final_pdf_path = Path(file_path).parent / f"word_{uuid.uuid4().hex[:8]}.pdf"
                    abs_output = str(final_pdf_path.absolute())
                    
                    doc = word.Documents.Open(abs_input, ReadOnly=True)
                    # wdFormatPDF = 17
                    doc.SaveAs(abs_output, FileFormat=17)
                    doc.Close(0) # wdDoNotSaveChanges = 0
                    word.Quit()
                    
                    if os.path.exists(abs_output):
                        print(f"DEBUG: Word conversion successful: {final_pdf_path.name}")
                        return str(final_pdf_path)
                finally:
                    pythoncom.CoUninitialize()
            except Exception as e:
                print(f"Microsoft Word conversion failed: {e}")

        # 3. Fallback: For DOCX specifically, try a basic rendering if others fail
        if file_type in ['docx', 'doc']:
            try:
                # Basic fidelity fallback (current implementation)
                # Note: true high-fidelity DOCX rendering without LibreOffice/Word 
                # requires complex specialized libraries.
                from docx import Document
                doc = Document(file_path)
                text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                # Create PDF from text with basic formatting
                pdf_path = Path(file_path).parent / f"fallback_{uuid.uuid4().hex[:8]}.pdf"
                c = canvas.Canvas(str(pdf_path), pagesize=A4)
                width, height = A4
                
                c.setFont("Helvetica-Bold", 14)
                c.drawString(50, height - 50, f"Document: {Path(file_path).name}")
                c.line(50, height - 60, width - 50, height - 60)
                
                y = height - 80
                lines = text_content.split('\n')
                for line in lines:
                    if y < 50:
                        c.showPage()
                        c.setFont("Helvetica", 10)
                        y = height - 50
                    
                    # Wrap text simple
                    wrapped_line = line[:100]
                    c.drawString(50, y, wrapped_line)
                    y -= 15
                
                c.save()
                return str(pdf_path)
            except Exception as e:
                print(f"Low-fidelity fallback failed: {e}")
        
        # Return original path if no conversion could be performed
        return file_path
        
    except Exception as e:
        print(f"Critical error in convert_to_pdf: {e}")
        return file_path

def combine_pages_into_layout_sheets(page_preview_paths, layout_pages, preview_dir):
    """
    LAYOUT FIX: Combine individual page previews into layout sheets.
    
    BUG FIX: Previously duplicated the last page when padding partial sheets.
    Now uses white placeholders for padding, ensuring no page duplication.
    Works identically for all document types: PDF, DOCX, DOC, ODT, PPTX, XLSX, images, etc.
    
    Takes a list of individual page preview image paths and combines them
    into sheets based on layout_pages (e.g., 2 pages per sheet, 4 per sheet).
    
    Algorithm:
    1. Calculate number of sheets needed: ceil(num_pages / layout_pages)
    2. For each sheet:
       a. Extract pages for this sheet (with bounds checking)
       b. Load page preview images
       c. If fewer pages than layout_pages, pad with WHITE placeholders (never duplicate)
       d. Combine pages into grid layout
    3. Return sheet preview paths
    
    Args:
        page_preview_paths: List of paths to individual page preview images
        layout_pages: Number of document pages per sheet (1, 2, 4, 6, 8, 9, 16)
        preview_dir: Directory to save combined sheet previews
    
    Returns:
        tuple: (number_of_sheets, list of sheet preview paths)
    """
    import math
    
    if layout_pages <= 1 or not page_preview_paths:
        # No layout needed, return individual pages as-is
        return len(page_preview_paths), page_preview_paths
    
    # Calculate number of sheets needed
    num_pages = len(page_preview_paths)
    num_sheets = math.ceil(num_pages / layout_pages)
    
    # Get grid dimensions for layout
    rows, cols = get_grid_for_layout(layout_pages)
    
    sheet_preview_paths = []
    
    # BUG FIX: Process each sheet with strict bounds checking
    # Calculate sheet ranges carefully to never exceed available pages
    for sheet_idx in range(num_sheets):
        # Get pages for this sheet with bounds checking
        start_page_idx = sheet_idx * layout_pages
        end_page_idx = min(start_page_idx + layout_pages, num_pages)  # Never exceed num_pages
        
        # Bounds check: ensure start_page_idx is valid
        if start_page_idx >= num_pages:
            print(f"LAYOUT FIX WARNING: Sheet {sheet_idx + 1} start index {start_page_idx} >= total pages {num_pages}, skipping")
            break
        
        # Extract page paths for this sheet (slice is safe, returns empty if out of bounds)
        sheet_page_paths = page_preview_paths[start_page_idx:end_page_idx]
        
        # Log for verification
        print(f"LAYOUT FIX: Processing sheet {sheet_idx + 1}/{num_sheets}: pages {start_page_idx + 1}-{end_page_idx} "
              f"({len(sheet_page_paths)} pages, need {layout_pages} per sheet)")
        
        # BUG FIX: Load page preview images with bounds checking
        # Never duplicate the last page - always use white placeholders for padding
        page_images = []
        for page_path in sheet_page_paths:
            try:
                img = Image.open(page_path)
                page_images.append(img)
            except Exception as e:
                print(f"Error loading page preview {page_path}: {e}")
                # Add white placeholder if page fails to load (never duplicate)
                placeholder = Image.new('RGB', (300, 400), 'white')
                page_images.append(placeholder)
        
        # BUG FIX: Pad with white placeholders only (never duplicate last page)
        # This ensures partial sheets show blank spaces instead of duplicate content
        # Works identically for all document types: PDF, DOCX, converted PDFs, images, etc.
        num_pages_in_sheet = len(page_images)
        if num_pages_in_sheet < layout_pages:
            # Get dimensions from first page for consistent placeholder size
            placeholder_size = (300, 400)  # Default
            if page_images:
                placeholder_size = (page_images[0].width, page_images[0].height)
            
            # Pad with white placeholders (never duplicate)
            pages_to_pad = layout_pages - num_pages_in_sheet
            for _ in range(pages_to_pad):
                white_placeholder = Image.new('RGB', placeholder_size, 'white')
                page_images.append(white_placeholder)
            
            # Log for verification (only in debug scenarios)
            if num_pages_in_sheet > 0:
                print(f"LAYOUT FIX: Sheet {sheet_idx + 1} padded with {pages_to_pad} white placeholder(s) "
                      f"(had {num_pages_in_sheet} pages, need {layout_pages} per sheet)")
        
        # Combine pages into layout sheet
        if layout_pages == 1:
            # Single page, no combination needed
            combined_img = page_images[0]
        else:
            # Use generic grid layout to combine pages
            # Start with first page to get dimensions
            base_img = page_images[0] if page_images else Image.new('RGB', (300, 400), 'white')
            
            # Calculate cell dimensions
            gap = 20
            target_width = min(800, base_img.width * cols)
            target_height = min(800, base_img.height * rows)
            cell_width = int((target_width - gap * (cols - 1)) / cols)
            cell_height = int((target_height - gap * (rows - 1)) / rows)
            
            # Preserve aspect ratio
            aspect = base_img.width / base_img.height
            if cell_width / cell_height > aspect:
                cell_width = int(cell_height * aspect)
            else:
                cell_height = int(cell_width / aspect)
            
            # Create combined sheet image
            combined_img = Image.new('RGB', 
                (cell_width * cols + gap * (cols - 1), 
                 cell_height * rows + gap * (rows - 1)), 
                'white')
            
            # BUG FIX: Place each page in the grid with strict bounds checking
            # Only place pages that actually exist - never index beyond available pages
            # This prevents any possibility of duplicating or reusing pages
            for idx in range(layout_pages):
                # Bounds check: ensure we don't go beyond available page_images
                if idx >= len(page_images):
                    # This should never happen after padding fix, but safety check
                    print(f"LAYOUT FIX WARNING: Index {idx} beyond page_images length {len(page_images)}, skipping")
                    break
                
                page_img = page_images[idx]
                r = idx // cols
                c = idx % cols
                
                # Resize page to fit cell
                resized = page_img.copy()
                resized.thumbnail((cell_width, cell_height), Image.Resampling.LANCZOS)
                
                # Calculate position
                x = c * (cell_width + gap)
                y = r * (cell_height + gap)
                
                # Center the image in the cell if it's smaller
                x_offset = (cell_width - resized.width) // 2
                y_offset = (cell_height - resized.height) // 2
                
                combined_img.paste(resized, (x + x_offset, y + y_offset))
        
        # Save combined sheet preview
        sheet_filename = f"preview_sheet_{uuid.uuid4().hex[:8]}_sheet_{sheet_idx + 1}.png"
        sheet_path = preview_dir / sheet_filename
        combined_img.save(sheet_path, "PNG", quality=95)
        sheet_preview_paths.append(str(sheet_path))
    
    return num_sheets, sheet_preview_paths

def generate_multi_page_previews(file_path, file_type, page_size="A4", orientation="Portrait", color_mode="Color", layout_pages=1):
    """
    Generate preview images for ALL pages of a document.
    Handles both local paths and Cloudinary URLs.
    """
    # ROOT FIX: Ensure local path
    local_path, is_temp = ensure_local_path(file_path)
    
    try:
        # Create preview directory
        if is_url(file_path):
            preview_dir = UPLOAD_FOLDER / "previews"
        else:
            preview_dir = Path(local_path).parent / "previews"
        preview_dir.mkdir(exist_ok=True)
        
        preview_paths = []
        total_pages = 0
        pdf_path = local_path
        
        # Step 1: Normalize non-PDF files for preview
        DOCUMENT_TYPES = ['docx', 'doc', 'pptx', 'xlsx', 'odt', 'ods', 'odp']
        if file_type.lower() in DOCUMENT_TYPES:
            normalized_path = normalize_document_for_preview(local_path, file_type)
            if normalized_path != local_path:
                local_path = normalized_path
                pdf_path = normalized_path  # ← ADD THIS LINE
                file_type = 'pdf'
        elif file_type == 'pdf':
            pdf_path = local_path
        

        # Step 2: Generate previews from PDF (or original if PDF)
        if file_type == 'pdf' or (pdf_path != local_path and Path(pdf_path).exists()):
            try:
                import fitz  # PyMuPDF
                
                pdf_document = fitz.open(pdf_path)
                total_pages = len(pdf_document)
                
                # Generate preview for each page
                for page_num in range(total_pages):
                    try:
                        page = pdf_document[page_num]
                        page_rect = page.rect
                        
                        # Calculate zoom based on page size
                        zoom_factor = get_zoom_for_page_size(page_size, page_rect)
                        
                        # Apply orientation transformation
                        if orientation == "Landscape":
                            if page_rect.width < page_rect.height:
                                mat = fitz.Matrix(0, 1, -1, 0, page_rect.height, 0)
                                mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
                            else:
                                mat = fitz.Matrix(zoom_factor, zoom_factor)
                        else:
                            if page_rect.width > page_rect.height:
                                mat = fitz.Matrix(0, -1, 1, 0, 0, page_rect.width)
                                mat = mat * fitz.Matrix(zoom_factor, zoom_factor)
                            else:
                                mat = fitz.Matrix(zoom_factor, zoom_factor)
                        
                        # Render page to image
                        pix = page.get_pixmap(matrix=mat, alpha=False)
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        
                        # Resize to preview size while maintaining aspect ratio
                        img.thumbnail((600, 800), Image.Resampling.LANCZOS)
                        
                        # Apply color mode conversion
                        img = apply_color_mode(img, color_mode)
                        
                        # Save preview image
                        preview_filename = f"preview_{uuid.uuid4().hex[:8]}_page_{page_num + 1}.png"
                        preview_path = preview_dir / preview_filename
                        img.save(preview_path, "PNG", quality=95)
                        preview_paths.append(str(preview_path))
                        
                    except Exception as e:
                        print(f"Error generating preview for page {page_num + 1}: {e}")
                        # Add fallback preview
                        fallback_path = preview_dir / f"preview_not_available_page_{page_num + 1}.png"
                        create_fallback_preview_page(fallback_path, page_num + 1, total_pages)
                        preview_paths.append(str(fallback_path))
                
                pdf_document.close()
                
                # Cleanup converted PDF ONLY if it was a one-off temporary file
                # If it's in the preview_cache, DO NOT delete it as it's meant to be reused
                if pdf_path != local_path and Path(pdf_path).exists() and "preview_cache" not in str(pdf_path):
                    try:
                        os.remove(pdf_path)
                    except:
                        pass
                
            except ImportError:
                # PyMuPDF not available, fallback to PyPDF2 (limited)
                try:
                    with open(pdf_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        total_pages = len(pdf_reader.pages)
                        
                        # Generate simple previews for each page
                        for page_num in range(total_pages):
                            preview_filename = f"preview_{uuid.uuid4().hex[:8]}_page_{page_num + 1}.png"
                            preview_path = preview_dir / preview_filename
                            create_fallback_preview_page(preview_path, page_num + 1, total_pages)
                            preview_paths.append(str(preview_path))
                except Exception as e:
                    print(f"Error with PyPDF2 fallback: {e}")
                    return (0, [])
        
        # Step 3: Handle images (including multi-frame TIFF)
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
            try:
                img = Image.open(local_path)
                
                # Handle multi-frame TIFF
                if file_type == 'tiff':
                    frame_count = 0
                    try:
                        while True:
                            # Apply orientation
                            frame_img = img.copy()
                            if orientation == "Landscape" and frame_img.width < frame_img.height:
                                frame_img = frame_img.rotate(90, expand=True)
                            elif orientation == "Portrait" and frame_img.width > frame_img.height:
                                frame_img = frame_img.rotate(-90, expand=True)
                            
                            # Resize
                            frame_img.thumbnail((600, 800), Image.Resampling.LANCZOS)
                            
                            # Apply color mode
                            frame_img = apply_color_mode(frame_img, color_mode)
                            
                            # Save preview
                            preview_filename = f"preview_{uuid.uuid4().hex[:8]}_page_{frame_count + 1}.png"
                            preview_path = preview_dir / preview_filename
                            frame_img.save(preview_path, "PNG", quality=95)
                            preview_paths.append(str(preview_path))
                            
                            frame_count += 1
                            try:
                                img.seek(frame_count)
                            except EOFError:
                                break
                        
                        # Set total pages after processing all frames
                        total_pages = frame_count
                    except Exception as e:
                        print(f"Error processing TIFF frames: {e}")
                        # Fallback: treat as single image
                        if not preview_paths:
                            preview_filename = f"preview_{uuid.uuid4().hex[:8]}_page_1.png"
                            preview_path = preview_dir / preview_filename
                            frame_img = apply_color_mode(img, color_mode)
                            frame_img.thumbnail((600, 800), Image.Resampling.LANCZOS)
                            frame_img.save(preview_path, "PNG", quality=95)
                            preview_paths.append(str(preview_path))
                            total_pages = 1
                else:
                    # Single image file
                    if img.mode in ('RGBA', 'LA', 'P'):
                        img = img.convert('RGB')
                    
                    # Apply orientation
                    if orientation == "Landscape" and img.width < img.height:
                        img = img.rotate(90, expand=True)
                    elif orientation == "Portrait" and img.width > img.height:
                        img = img.rotate(-90, expand=True)
                    
                    # Resize
                    img.thumbnail((600, 800), Image.Resampling.LANCZOS)
                    
                    # Apply color mode
                    img = apply_color_mode(img, color_mode)
                    
                    # Save preview
                    preview_filename = f"preview_{uuid.uuid4().hex[:8]}_page_1.png"
                    preview_path = preview_dir / preview_filename
                    img.save(preview_path, "PNG", quality=95)
                    preview_paths.append(str(preview_path))
                    total_pages = 1
                    
            except Exception as e:
                print(f"Error processing image: {e}")
                return (0, [])
        
        else:
            # Unsupported file type - create single fallback preview
            preview_filename = f"preview_{uuid.uuid4().hex[:8]}_page_1.png"
            preview_path = preview_dir / preview_filename
            create_fallback_preview_page(preview_path, 1, 1, file_type=file_type)
            preview_paths.append(str(preview_path))
            total_pages = 1
        
        # LAYOUT FIX: Combine individual page previews into layout sheets if layout_pages > 1
        if layout_pages > 1 and preview_paths:
            num_sheets, sheet_preview_paths = combine_pages_into_layout_sheets(
                preview_paths, 
                layout_pages, 
                preview_dir
            )
            # Return sheet previews instead of individual page previews
            return (num_sheets, sheet_preview_paths)
        
        # Return individual page previews (layout_pages == 1)
        return (total_pages, preview_paths)
        
    except Exception as e:
        print(f"Error generating multi-page previews: {e}")
        # Return fallback preview
        try:
            preview_dir = Path(file_path).parent / "previews"
            preview_dir.mkdir(exist_ok=True)
            preview_filename = f"preview_not_available.png"
            preview_path = preview_dir / preview_filename
            create_fallback_preview_page(preview_path, 1, 1, error=str(e))
            return (1, [str(preview_path)])
        except:
            return (0, [])
    finally:
        if is_temp and os.path.exists(local_path):
            try: os.remove(local_path)
            except: pass

def create_fallback_preview_page(preview_path, page_num, total_pages, file_type=None, error=None):
    """
    Create a fallback preview image when page generation fails.
    
    Args:
        preview_path: Path to save the preview
        page_num: Current page number
        total_pages: Total number of pages
        file_type: Optional file type for display
        error: Optional error message
    """
    try:
        from PIL import ImageDraw, ImageFont
        
        img = Image.new('RGB', (600, 800), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("arial.ttf", 20)
            small_font = ImageFont.truetype("arial.ttf", 14)
        except:
            try:
                font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 20)
                small_font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 14)
            except:
                font = ImageFont.load_default()
                small_font = ImageFont.load_default()
        
        y = 50
        draw.text((50, y), "Preview Not Available", fill='gray', font=font)
        y += 40
        draw.text((50, y), f"Page {page_num} of {total_pages}", fill='black', font=small_font)
        
        if file_type:
            y += 30
            draw.text((50, y), f"Type: {file_type.upper()}", fill='gray', font=small_font)
        
        if error:
            y += 30
            draw.text((50, y), f"Error: {str(error)[:60]}", fill='red', font=small_font)
        
        img.save(preview_path, "PNG")
    except Exception as e:
        print(f"Error creating fallback preview: {e}")

def classify_color_pages(file_path, file_type):
    """
    Classify each page of a document as 'color' or 'bw'.
    Handles both local paths and Cloudinary URLs.
    """
    # ROOT FIX: Ensure local path
    local_path, is_temp = ensure_local_path(file_path)
    
    try:
        import fitz
        
        # Step 0: Normalize document to PDF if needed
        if file_type != 'pdf':
            normalized_path = normalize_document_for_preview(local_path, file_type)
            if normalized_path != local_path:
                local_path = normalized_path
                file_type = 'pdf'
            
        doc = fitz.open(local_path)
        total_pages = len(doc)
        results = {}
        
        for i in range(total_pages):
            page = doc[i]
            pix = page.get_pixmap(dpi=300)
            
            # Get samples as bytes
            samples_bytes = pix.samples
            width = pix.width
            height = pix.height
            n = pix.n # components per pixel (usually 3 for RGB)
            
            is_color = False
            tr = 15
            
            # Phase 1: Fast Coarse Scan (Preserves performance for large images)
            step = 10
            colored_pixels = 0
            for y in range(0, height, step):
                for x in range(0, width, step):
                    pos = (y * width + x) * n
                    r, g, b = samples_bytes[pos:pos+3]
                    if abs(int(r) - int(g)) > tr or abs(int(g) - int(b)) > tr or abs(int(r) - int(b)) > tr:
                        colored_pixels += 1
                        if colored_pixels >= 10: # Early exit: clear color content found
                            is_color = True
                            break
                if is_color: break
            
            # Phase 2: Fallback Fine Scan (NEW: Only if Phase 1 fails to find small elements)
            # Higher resolution scan to catch small logos, text, or thin colored lines
            if not is_color:
                step = 2
                colored_pixels = 0
                for y in range(0, height, step):
                    for x in range(0, width, step):
                        pos = (y * width + x) * n
                        r, g, b = samples_bytes[pos:pos+3]
                        if abs(int(r) - int(g)) > tr or abs(int(g) - int(b)) > tr or abs(int(r) - int(b)) > tr:
                            colored_pixels += 1
                            if colored_pixels >= 5: # Critical threshold: existence of small color confirmed
                                is_color = True
                                break
                    if is_color: break
            
            results[i + 1] = "color" if is_color else "bw"
                
        doc.close()
        return results
    except Exception as e:
        print(f"Error classifying color pages: {e}")
        return None
    finally:
        if is_temp and os.path.exists(local_path):
            try: os.remove(local_path)
            except: pass

def build_color_page_dict(color_pages_list, total_pages):
    """Build a color_page_dict from user-selected color page numbers.

    Converts a list of 1-indexed page numbers (user-selected as "color") into
    the {page_number: 'color'|'bw'} dict consumed by calculate_billing().

    Args:
        color_pages_list: list/set of 1-indexed page numbers to treat as color.
                          None or empty → returns None (all-color fallback).
        total_pages: total document page count.

    Returns:
        dict or None: {1: 'color', 2: 'bw', ...} or None if no selection.

    Safety:
        - Returns None on any error (backward-compatible: all-color billing).
        - Ignores out-of-range page numbers silently.
        - Pure function — no I/O, no side effects.
    """
    try:
        if not color_pages_list or not total_pages or total_pages < 1:
            return None
        color_set = set(int(p) for p in color_pages_list if isinstance(p, (int, float)) and p > 0)
        if not color_set:
            return None
        return {pg: ('color' if pg in color_set else 'bw') for pg in range(1, total_pages + 1)}
    except Exception:
        return None


def calculate_billing(color_mode, print_side, copies, layout_pages, selected_pages, color_page_dict, pricing):
    """
    Shared billing calculation logic for both PRINT and XEROX flows.
    """
    import math
    is_double = print_side.lower() in ['double', 'duplex']
    
    bw_single = pricing.get('bw_single', 2.0)
    bw_double = pricing.get('bw_double', 1.5)
    color_single = pricing.get('color_single', 10.0)
    color_double = pricing.get('color_double', 8.0)
    
    rate_color = color_double if is_double else color_single
    rate_bw = bw_double if is_double else bw_single
    
    color_sheets = 0
    bw_sheets = 0
    
    if layout_pages > 1:
        for i in range(0, len(selected_pages), layout_pages):
            batch = selected_pages[i:i + layout_pages]
            is_sheet_color = False
            
            if color_mode.lower() != 'black & white':
                if color_page_dict is None:
                    is_sheet_color = True
                else:
                    for p in batch:
                        if color_page_dict.get(p) == "color":
                            is_sheet_color = True
                            break
            
            if is_sheet_color:
                color_sheets += 1
            else:
                bw_sheets += 1
    else:
        for p in selected_pages:
            is_page_color = False
            if color_mode.lower() != 'black & white':
                if color_page_dict is None or color_page_dict.get(p) == "color":
                    is_page_color = True
            
            if is_page_color:
                color_sheets += 1
            else:
                bw_sheets += 1
    # Duplex Sheet Normalization Rule
    if is_double:
        import math
        bw_sheets = math.ceil(bw_sheets / 2)
        color_sheets = math.ceil(color_sheets / 2)
                
    total_amount = (color_sheets * rate_color + bw_sheets * rate_bw) * copies
    # CEILING ROUNDING: Round up to next whole rupee (e.g. 4.5 -> 5, 10.1 -> 11, 10 -> 10)
    total_amount = math.ceil(total_amount)
    return {
        'total_amount': total_amount,
        'color_sheets': color_sheets,
        'bw_sheets': bw_sheets,
        'page_count': color_sheets + bw_sheets
    }

def cleanup_temp_files(file_path):
    """
    Clean up temporary files
    
    Args:
        file_path (str): Path to file to delete
    """
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        print(f"Error cleaning up file {file_path}: {e}")

def combine_images_to_pdf(image_paths, output_pdf_path):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.utils import ImageReader
        from PIL import ImageOps

        c = canvas.Canvas(output_pdf_path, pagesize=A4)

        for image_path in image_paths:
            try:
                img = Image.open(image_path)
                img = ImageOps.exif_transpose(img)  # Fix EXIF rotation
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                img_width, img_height = img.size

                # Fix orientation — match page to image
                if img_width > img_height:
                    page_w, page_h = A4[1], A4[0]  # landscape
                else:
                    page_w, page_h = A4              # portrait

                c.setPageSize((page_w, page_h))  # per page apply

                scale = min(page_w / img_width, page_h / img_height)
                scaled_w = img_width * scale
                scaled_h = img_height * scale
                x = (page_w - scaled_w) / 2
                y = (page_h - scaled_h) / 2

                img_bytes = io.BytesIO()
                img.save(img_bytes, format='PNG')
                img_bytes.seek(0)

                c.drawImage(ImageReader(img_bytes), x, y,
                            width=scaled_w, height=scaled_h)
                c.showPage()

            except Exception as e:
                logger.error(f"Error processing image {image_path}: {e}")
                continue

        c.save()
        return output_pdf_path

    except Exception as e:
        logger.error(f"Error creating PDF from images: {e}")
        raise

def cleanup_old_uploads(max_age_hours: int = 72):
    """Delete upload and preview files older than max_age_hours (best-effort)."""
    cutoff = time.time() - max_age_hours * 3600
    try:
        for root, dirs, files in os.walk(UPLOAD_FOLDER):
            for name in files:
                path = os.path.join(root, name)
                try:
                    if os.path.getmtime(path) < cutoff:
                        os.remove(path)
                except Exception:
                    pass
    except Exception:
        pass

import hashlib


def _compute_proxy_lookup_key(original_filename, file_size):
    """Compute a deterministic cache key for proxy file lookup.

    Uses ONLY metadata (original filename + file size) — no file I/O,
    no image processing, no content reading.  The same inputs always
    produce the same key, enabling the preview resolver to find a proxy
    that was generated during upload without any shared state or DB.

    Args:
        original_filename (str): Browser-provided original filename.
        file_size (int): File size in bytes.

    Returns:
        str: 12-character hex digest, or None on error.
    """
    try:
        raw = f"{original_filename}:{file_size}".encode('utf-8')
        return hashlib.md5(raw).hexdigest()[:12]
    except Exception:
        return None


def resolve_preview_source(original_path, file_type, original_filename=None):
    """Pure cache lookup for a pre-generated proxy asset.

    This function does NO image processing, NO proxy generation, NO file
    reading.  It only:
      1. Computes a deterministic cache key from metadata
      2. Checks whether a matching proxy JPEG exists on disk
      3. Returns the proxy path (cache hit) or original path (cache miss)

    The proxy was already generated asynchronously by Phase 1B during upload.
    This function merely CONSUMES that cached asset.

    PREVIEW-ONLY:  Never called by the print pipeline.

    Args:
        original_path (str): Path to the temp file saved by the preview route.
        file_type (str): File extension without dot.
        original_filename (str, optional): Browser-provided original filename.
            Used together with file size to compute the cache key.
            Falls back to basename of original_path if not provided.

    Returns:
        tuple: (resolved_path: str, is_proxy: bool)
    """
    PROXY_IMAGE_TYPES = {'jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff'}
    ft = (file_type or '').lower().strip()

    # Only look up proxies for image types (PDFs/docs need multi-page fidelity)
    if ft not in PROXY_IMAGE_TYPES:
        logger.debug(f"PREVIEW SOURCE: Using original (type={ft})")
        return original_path, False

    if not original_path or not os.path.exists(original_path):
        logger.debug("PREVIEW SOURCE: Using original (path missing)")
        return original_path, False

    # Compute deterministic lookup key — metadata only, no file content read
    try:
        file_size = os.path.getsize(original_path)
    except Exception:
        return original_path, False

    fname = original_filename or os.path.basename(original_path)
    lookup_key = _compute_proxy_lookup_key(fname, file_size)
    if not lookup_key:
        return original_path, False

    # Check if Phase 1B already generated a proxy with this key
    proxy_path = UPLOAD_FOLDER / "proxies" / f"proxy_{lookup_key}.jpg"
    if proxy_path.exists():
        logger.debug(f"PREVIEW SOURCE: Cache HIT — key={lookup_key}")
        return str(proxy_path), True

    logger.debug(f"PREVIEW SOURCE: Cache MISS — key={lookup_key}, using original")
    return original_path, False


def generate_proxy_asset(file_path, file_type, proxy_id=None):
    """
    Generate a lightweight proxy JPEG from an uploaded file for fast UI previews.
    
    This is an ISOLATED utility function — it does NOT modify any existing
    preview, upload, or print pipeline. It is designed to be called independently
    and produces a small (~100-200KB) JPEG proxy suitable for rapid UI rendering.
    
    The ORIGINAL file is NEVER modified. The proxy is a separate, disposable asset.
    
    Supported inputs:
        - Images (jpg/jpeg/png/webp/bmp/tiff): Thumbnail with EXIF correction
        - PDF: First page rendered at 0.5x scale
        - Documents (docx/doc/pptx/xlsx/odt/ods/odp): Converted to PDF first,
          then first page rendered (reuses existing convert_to_pdf)
    
    Args:
        file_path (str): Absolute path to the source file (local path, NOT a URL)
        file_type (str): File extension without dot (e.g. 'jpg', 'pdf', 'docx')
        proxy_id (str, optional): Unique identifier for the proxy file.
                                  If None, a UUID hex prefix is generated.
    
    Returns:
        dict: {
            'success': bool,
            'proxy_id': str,          # Unique proxy identifier
            'proxy_path': str,        # Absolute path to generated proxy JPEG
            'source_type': str,       # 'image', 'pdf', or 'document'
            'original_path': str,     # Original file path (unchanged)
            'proxy_size_bytes': int,  # Size of generated proxy file
        }
        On failure: { 'success': False, 'error': str, 'original_path': str }
    
    Memory budget:
        - Image (20MB JPEG): ~38MB peak (PIL decode + thumbnail)
        - PDF: ~5MB peak (fitz at 0.5x scale)
        - Document: ~5MB peak (after convert_to_pdf)
    
    Safety:
        - Never modifies the source file
        - Never touches print pipeline functions
        - Never touches existing preview functions
        - Fully rollback-safe (delete the function to remove)
    """
    from PIL import ImageOps
    
    # Validate inputs
    if not file_path or not os.path.exists(file_path):
        return {
            'success': False,
            'error': f'Source file not found: {file_path}',
            'original_path': file_path
        }
    
    file_type = (file_type or '').lower().strip().lstrip('.')
    if not file_type:
        return {
            'success': False,
            'error': 'file_type is required',
            'original_path': file_path
        }
    
    # Generate proxy ID if not provided
    if not proxy_id:
        proxy_id = uuid.uuid4().hex[:12]
    
    # Create proxy output directory: uploads/proxies/
    proxy_dir = UPLOAD_FOLDER / "proxies"
    try:
        proxy_dir.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        return {
            'success': False,
            'error': f'Failed to create proxy directory: {e}',
            'original_path': file_path
        }
    
    proxy_filename = f"proxy_{proxy_id}.jpg"
    proxy_path = proxy_dir / proxy_filename
    
    # Classify input type
    IMAGE_TYPES = {'jpg', 'jpeg', 'png', 'webp', 'bmp', 'tiff'}
    PDF_TYPES = {'pdf'}
    DOCUMENT_TYPES = {'docx', 'doc', 'pptx', 'ppt', 'xlsx', 'xls', 'odt', 'ods', 'odp'}
    
    try:
        # ── BRANCH 1: Image files ──────────────────────────────────────
        if file_type in IMAGE_TYPES:
            img = Image.open(file_path)
            
            # CRITICAL: Fix mobile camera EXIF orientation before thumbnailing.
            # Without this, portrait photos from phones appear rotated.
            img = ImageOps.exif_transpose(img)
            
            # Convert to RGB if needed (handles RGBA, LA, P, CMYK modes)
            if img.mode not in ('RGB',):
                img = img.convert('RGB')
            
            # Thumbnail to max 1024px on longest side (preserves aspect ratio)
            img.thumbnail((1024, 1024), Image.Resampling.LANCZOS)
            
            # Save optimized JPEG proxy
            img.save(str(proxy_path), 'JPEG', quality=70, optimize=True)
            
            proxy_size = os.path.getsize(str(proxy_path))
            logger.info(
                f"PROXY GEN [image]: {Path(file_path).name} → {proxy_filename} "
                f"({proxy_size / 1024:.0f}KB, {img.size[0]}x{img.size[1]})"
            )
            
            return {
                'success': True,
                'proxy_id': proxy_id,
                'proxy_path': str(proxy_path),
                'source_type': 'image',
                'original_path': file_path,
                'proxy_size_bytes': proxy_size
            }
        
        # ── BRANCH 2: PDF files ────────────────────────────────────────
        elif file_type in PDF_TYPES:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            if len(doc) == 0:
                doc.close()
                return {
                    'success': False,
                    'error': 'PDF has no pages',
                    'original_path': file_path
                }
            
            # Render FIRST page only at 0.5x scale (lightweight, low memory)
            page = doc[0]
            mat = fitz.Matrix(0.5, 0.5)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes("png")
            total_pdf_pages = len(doc)
            doc.close()
            
            # Load into PIL for consistent output processing
            img = Image.open(io.BytesIO(img_data))
            if img.mode not in ('RGB',):
                img = img.convert('RGB')
            
            # Scale up to 1024px width if the 0.5x render is too small
            # (A4 at 72 DPI × 0.5 = ~297px width — too small for preview)
            if img.width < 1024:
                ratio = 1024 / img.width
                new_width = 1024
                new_height = int(img.height * ratio)
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # Save optimized JPEG proxy
            img.save(str(proxy_path), 'JPEG', quality=70, optimize=True)
            
            proxy_size = os.path.getsize(str(proxy_path))
            logger.info(
                f"PROXY GEN [pdf]: {Path(file_path).name} → {proxy_filename} "
                f"({proxy_size / 1024:.0f}KB, {img.size[0]}x{img.size[1]}, page 1 of {total_pdf_pages})"
            )
            
            return {
                'success': True,
                'proxy_id': proxy_id,
                'proxy_path': str(proxy_path),
                'source_type': 'pdf',
                'original_path': file_path,
                'proxy_size_bytes': proxy_size
            }
        
        # ── BRANCH 3: Document files (DOCX/DOC/PPTX etc.) ─────────────
        elif file_type in DOCUMENT_TYPES:
            # Reuse existing convert_to_pdf() to get a PDF, then use PDF branch
            converted_pdf_path = convert_to_pdf(file_path, file_type)
            
            if converted_pdf_path == file_path:
                # convert_to_pdf returns original path on failure
                return {
                    'success': False,
                    'error': f'Failed to convert {file_type.upper()} to PDF for proxy generation',
                    'original_path': file_path
                }
            
            # Recursively call self with the converted PDF
            result = generate_proxy_asset(converted_pdf_path, 'pdf', proxy_id=proxy_id)
            
            # Update source_type to reflect original document type
            if result.get('success'):
                result['source_type'] = 'document'
                result['original_path'] = file_path  # Point back to original, not converted PDF
            
            # Clean up the temporary converted PDF
            try:
                if converted_pdf_path != file_path and os.path.exists(converted_pdf_path):
                    os.remove(converted_pdf_path)
            except Exception:
                pass  # Non-fatal: temp file cleanup failure
            
            return result
        
        # ── UNSUPPORTED TYPE ───────────────────────────────────────────
        else:
            return {
                'success': False,
                'error': f'Unsupported file type for proxy generation: {file_type}',
                'original_path': file_path
            }
    
    except Exception as e:
        logger.error(f"PROXY GEN ERROR: {file_type} file '{Path(file_path).name}': {e}")
        # Clean up partial proxy file on error
        try:
            if proxy_path.exists():
                os.remove(str(proxy_path))
        except Exception:
            pass
        
        return {
            'success': False,
            'error': f'Proxy generation failed: {str(e)}',
            'original_path': file_path
        }
